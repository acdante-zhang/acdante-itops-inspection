#!/usr/bin/env python3
"""
Acdante ITOps - 综合巡检报告Word模板生成器
参考义乌社保巡检报告风格，生成标准化巡检模板
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ============================================================
# 全局样式设置
# ============================================================
FONT_NAME = '宋体'
FONT_NAME_TITLE = '黑体'
FONT_COLOR = RGBColor(0, 0, 0)
FONT_SIZE = Pt(10.5)
FONT_SIZE_SMALL = Pt(9)
FONT_SIZE_TITLE = Pt(22)
FONT_SIZE_H1 = Pt(16)
FONT_SIZE_H2 = Pt(14)

style = doc.styles['Normal']
font = style.font
font.name = FONT_NAME
font.color.rgb = FONT_COLOR
font.size = FONT_SIZE
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

# ============================================================
# 辅助函数
# ============================================================
def set_cell(cell, text, bold=False, size=None, align='left', bg=None, font_name=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.font.name = font_name or FONT_NAME
    run.font.color.rgb = FONT_COLOR
    run.font.size = size or FONT_SIZE
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name or FONT_NAME)
    if bg:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tcBorders>')
    tcPr.append(tcBorders)

def merge_cells(table, row1, col1, row2, col2):
    cell1 = table.cell(row1, col1)
    cell2 = table.cell(row2, col2)
    cell1.merge(cell2)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_NAME_TITLE
        run.font.color.rgb = FONT_COLOR
        run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME_TITLE)
    return h

def add_normal(text, bold=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.color.rgb = FONT_COLOR
    run.font.size = size or FONT_SIZE
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    return p

def add_status_tag(cell, status='normal'):
    """添加状态标注：normal=正常 √, abnormal=异常 ×"""
    if status == 'normal':
        set_cell(cell, '√ 正常', bold=True, size=FONT_SIZE_SMALL, align='center', bg='D5F5E3')
    elif status == 'abnormal':
        set_cell(cell, '× 异常', bold=True, size=FONT_SIZE_SMALL, align='center', bg='FADBD8')
    else:
        set_cell(cell, status, size=FONT_SIZE_SMALL, align='center')

# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('【客户名称】')
run.font.name = FONT_NAME_TITLE
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0, 51, 102)
run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME_TITLE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('【系统名称】')
run.font.name = FONT_NAME_TITLE
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0, 51, 102)
run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME_TITLE)

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026年M【XX】月')
run.font.name = FONT_NAME_TITLE
run.font.size = Pt(16)
run.font.color.rgb = FONT_COLOR
run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME_TITLE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('度')
run.font.name = FONT_NAME_TITLE
run.font.size = Pt(16)
run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME_TITLE)

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('巡检服务报告')
run.font.name = FONT_NAME_TITLE
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)
run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME_TITLE)

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('【运维服务商名称】')
run.font.name = FONT_NAME
run.font.size = Pt(14)
run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

doc.add_page_break()

# ============================================================
# 文档控制
# ============================================================
add_heading_styled('文档控制', level=2)

# 客户联系人
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
set_cell(table.rows[0].cells[0], '客户联系人', bold=True, bg='D6EAF8')
set_cell(table.rows[0].cells[1], '【联系人姓名】')
set_cell(table.rows[0].cells[2], '【联系人姓名】')
set_cell(table.rows[0].cells[3], '【联系人姓名】')
set_cell(table.rows[1].cells[0], '巡检报告编写', bold=True, bg='D6EAF8')
set_cell(table.rows[1].cells[1], '【编写人】/【编写人】')
set_cell(table.rows[1].cells[2], '编写日期')
set_cell(table.rows[1].cells[3], '【YYYY.MM.DD】')
set_cell(table.rows[2].cells[0], '巡检报告审批', bold=True, bg='D6EAF8')
set_cell(table.rows[2].cells[1], '【审批人】')
set_cell(table.rows[2].cells[2], '审批日期')
set_cell(table.rows[2].cells[3], '【YYYY.MM.DD】')
set_cell(table.rows[3].cells[0], '巡检报告分发', bold=True, bg='D6EAF8')
set_cell(table.rows[3].cells[1], '【分发人】')
set_cell(table.rows[3].cells[2], '分发日期')
set_cell(table.rows[3].cells[3], '【YYYY.MM.DD】')

doc.add_paragraph()

# 版本修改记录
add_normal('版本修改记录', bold=True)
table = doc.add_table(rows=3, cols=4)
table.style = 'Table Grid'
headers = ['日期', '作者', '版本', '修改记录']
for i, h in enumerate(headers):
    set_cell(table.rows[0].cells[i], h, bold=True, bg='D6EAF8', align='center')
set_cell(table.rows[1].cells[0], '【YYYY.MM.DD】')
set_cell(table.rows[1].cells[1], '【作者】')
set_cell(table.rows[1].cells[2], 'V1.0')
set_cell(table.rows[1].cells[3], '创建文档')

doc.add_paragraph()

# 审阅记录
add_normal('审阅记录', bold=True)
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'
set_cell(table.rows[0].cells[0], '审阅人', bold=True, bg='D6EAF8', align='center')
set_cell(table.rows[0].cells[1], '备注', bold=True, bg='D6EAF8', align='center')
set_cell(table.rows[1].cells[0], '【审阅人】')
set_cell(table.rows[1].cells[1], '服务质量监督')
set_cell(table.rows[2].cells[0], '【审阅人】')
set_cell(table.rows[2].cells[1], '客户经理')

doc.add_page_break()

# ============================================================
# 目录
# ============================================================
add_heading_styled('目录', level=1)
add_normal('（生成文档后更新目录域）')
doc.add_page_break()

# ============================================================
# 第一章：巡检服务报告
# ============================================================
add_heading_styled('巡检服务报告', level=1)

# 巡检任务信息表
table = doc.add_table(rows=12, cols=6)
table.style = 'Table Grid'

# 标题行
merge_cells(table, 0, 0, 0, 5)
set_cell(table.rows[0].cells[0], '巡检任务信息', bold=True, size=Pt(12), align='center', bg='2E86C1')
table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

# 基本信息
set_cell(table.rows[1].cells[0], '用户单位', bold=True, bg='D6EAF8')
merge_cells(table, 1, 1, 1, 5)
set_cell(table.rows[1].cells[1], '【用户单位名称】')

set_cell(table.rows[2].cells[0], '巡检时间', bold=True, bg='D6EAF8')
set_cell(table.rows[2].cells[1], '【YYYY.MM.DD】')
set_cell(table.rows[2].cells[2], '')
set_cell(table.rows[2].cells[3], '巡检周期', bold=True, bg='D6EAF8')
merge_cells(table, 2, 4, 2, 5)
set_cell(table.rows[2].cells[4], '【月度/季度/年度】')

set_cell(table.rows[3].cells[0], '运维服务团队', bold=True, bg='D6EAF8')
merge_cells(table, 3, 1, 3, 5)
set_cell(table.rows[3].cells[1], '【运维团队名称】')

# 巡检工程师
set_cell(table.rows[4].cells[0], '巡检工程师', bold=True, bg='D6EAF8')
set_cell(table.rows[4].cells[1], '职责', bold=True, bg='D6EAF8', align='center')
merge_cells(table, 4, 2, 4, 3)
set_cell(table.rows[4].cells[2], '姓名', bold=True, bg='D6EAF8', align='center')
merge_cells(table, 4, 4, 4, 5)
set_cell(table.rows[4].cells[4], '联系方式', bold=True, bg='D6EAF8', align='center')

set_cell(table.rows[5].cells[0], '')
set_cell(table.rows[5].cells[1], '网络MA')
merge_cells(table, 5, 2, 5, 3)
set_cell(table.rows[5].cells[2], '【姓名】')
merge_cells(table, 5, 4, 5, 5)
set_cell(table.rows[5].cells[4], '【电话】')

set_cell(table.rows[6].cells[0], '')
set_cell(table.rows[6].cells[1], '主机和数据库MA')
merge_cells(table, 6, 2, 6, 3)
set_cell(table.rows[6].cells[2], '【姓名】')
merge_cells(table, 6, 4, 6, 5)
set_cell(table.rows[6].cells[4], '【电话】')

set_cell(table.rows[7].cells[0], '')
set_cell(table.rows[7].cells[1], '存储MA')
merge_cells(table, 7, 2, 7, 3)
set_cell(table.rows[7].cells[2], '【姓名】')
merge_cells(table, 7, 4, 7, 5)
set_cell(table.rows[7].cells[4], '【电话】')

# 巡检服务内容
merge_cells(table, 8, 0, 8, 5)
set_cell(table.rows[8].cells[0], '巡检服务内容', bold=True, bg='2E86C1')
table.rows[8].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

merge_cells(table, 9, 0, 9, 5)
set_cell(table.rows[9].cells[0], '根据维保合同设备清单，按照月度巡检计划到达客户现场，进行网络、服务器、存储硬件系统的巡检服务，及时发现设备运行中的问题和隐患，确保设备安全稳定运行。')

# 巡检结果及维护建议
merge_cells(table, 10, 0, 10, 5)
set_cell(table.rows[10].cells[0], '巡检结果及维护建议', bold=True, bg='2E86C1')
table.rows[10].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

merge_cells(table, 11, 0, 11, 5)
set_cell(table.rows[11].cells[0], '本次检查的总结：\n\n上次巡检主要问题总结：\n\n本次巡检主要问题总结：\n\n维护建议：')

doc.add_paragraph()

# ============================================================
# 第二章：设备清单总览
# ============================================================
add_heading_styled('一、硬件系统巡检清单', level=2)
add_normal('问题和概要见下表总结：')

# 设备清单表
table = doc.add_table(rows=21, cols=8)
table.style = 'Table Grid'
headers = ['序号', '类型', '设备类型', '品牌型号', '数量', '序列号', 'IP地址', '备注']
for i, h in enumerate(headers):
    set_cell(table.rows[0].cells[i], h, bold=True, bg='2E86C1', align='center', size=FONT_SIZE_SMALL)
    table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

# 示例设备数据
devices = [
    ['1', '网络', '核心路由器', 'H3C SR6608-X', '2', '【序列号】', '【IP】', '主/备'],
    ['2', '网络', '核心交换机', 'H3C S10508-V', '2', '【序列号】', '【IP】', '主/备'],
    ['3', '网络', '汇聚交换机', 'H3C S6800', '4', '【序列号】', '【IP】', ''],
    ['4', '网络', '防火墙', '【品牌型号】', '2', '【序列号】', '【IP】', '主/备'],
    ['5', '网络', '负载均衡', 'F5 BIG-IP', '2', '【序列号】', '【IP】', '主/备'],
    ['6', '网络', 'IPS/IDS', '【品牌型号】', '1', '【序列号】', '【IP】', ''],
    ['7', '网络', 'VPN设备', '【品牌型号】', '1', '【序列号】', '【IP】', ''],
    ['8', '网络', '堡垒机', '【品牌型号】', '1', '【序列号】', '【IP】', ''],
    ['9', '网络', 'WAF防火墙', '【品牌型号】', '1', '【序列号】', '【IP】', ''],
    ['10', '网络', '网闸', '【品牌型号】', '1', '【序列号】', '【IP】', '内/外'],
    ['11', '服务器', '刀片服务器', 'HP UIS8000', '3', '【序列号】', '【IP】', '生产/容灾'],
    ['12', '服务器', '小型机', 'IBM P750/S822', '3', '【序列号】', '【IP】', '数据库节点'],
    ['13', '服务器', 'x86服务器', '【品牌型号】', 'N', '【序列号】', '【IP】', '应用/Web'],
    ['14', '存储', '磁盘阵列', 'EMC VNX5600', '4', '【序列号】', '【IP】', '生产/容灾/测试'],
    ['15', '存储', '存储虚拟化', 'EMC VPLEX', '2', '【序列号】', '【IP】', '生产/容灾'],
    ['16', '存储', '光纤交换机', 'EMC 6505B', '6', '【序列号】', '【IP】', '生产/容灾/云区'],
    ['17', '存储', '磁带库', '昆腾Scalar i50', '1', '【序列号】', '【IP】', '备份'],
    ['18', '虚拟化', 'VMware集群', 'vSphere', '2', '-', '【IP】', '生产/测试'],
    ['19', '备份', '备份服务器', 'NetBackup', '1', '【序列号】', '【IP】', ''],
    ['20', '数据库', 'Oracle RAC', 'Oracle 19c', '1', '-', '【IP】', 'RAC集群'],
]

for r_idx, dev in enumerate(devices):
    for c_idx, val in enumerate(dev):
        set_cell(table.rows[r_idx + 1].cells[c_idx], val, size=FONT_SIZE_SMALL, align='center')

doc.add_page_break()

# ============================================================
# 第三章：巡检总结与问题剖析
# ============================================================
add_heading_styled('二、巡检总结', level=2)

add_normal('1. 巡检概况', bold=True)
add_normal('本次巡检共检查设备【XX】台，其中网络设备【XX】台、服务器【XX】台、存储设备【XX】台、安全设备【XX】台。巡检覆盖率达到100%。')

add_normal('2. 巡检结果统计', bold=True)

# 统计表
table = doc.add_table(rows=6, cols=5)
table.style = 'Table Grid'
set_cell(table.rows[0].cells[0], '设备类型', bold=True, bg='2E86C1', align='center')
set_cell(table.rows[0].cells[1], '检查数量', bold=True, bg='2E86C1', align='center')
set_cell(table.rows[0].cells[2], '正常', bold=True, bg='2E86C1', align='center')
set_cell(table.rows[0].cells[3], '异常', bold=True, bg='2E86C1', align='center')
set_cell(table.rows[0].cells[4], '健康度', bold=True, bg='2E86C1', align='center')
table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
table.rows[0].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
table.rows[0].cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
table.rows[0].cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
table.rows[0].cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

stats = [
    ['网络设备', '【N】', '【N】', '【N】', '【XX%】'],
    ['服务器', '【N】', '【N】', '【N】', '【XX%】'],
    ['存储设备', '【N】', '【N】', '【N】', '【XX%】'],
    ['安全设备', '【N】', '【N】', '【N】', '【XX%】'],
    ['合计', '【N】', '【N】', '【N】', '【XX%】'],
]
for r_idx, row_data in enumerate(stats):
    for c_idx, val in enumerate(row_data):
        bg = 'D5F5E3' if c_idx == 2 and r_idx < 5 else ('FADBD8' if c_idx == 3 and '【N】' != val else None)
        set_cell(table.rows[r_idx + 1].cells[c_idx], val, align='center', bg=bg)

doc.add_paragraph()

add_normal('3. 主要问题汇总', bold=True)

# 问题汇总表
table = doc.add_table(rows=6, cols=6)
table.style = 'Table Grid'
headers = ['序号', '设备名称', 'IP地址', '问题描述', '严重程度', '处理建议']
for i, h in enumerate(headers):
    set_cell(table.rows[0].cells[i], h, bold=True, bg='E74C3C', align='center', size=FONT_SIZE_SMALL)
    table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

for r in range(1, 6):
    set_cell(table.rows[r].cells[0], str(r), align='center', size=FONT_SIZE_SMALL)
    for c in range(1, 6):
        set_cell(table.rows[r].cells[c], '【填写】', size=FONT_SIZE_SMALL)

doc.add_paragraph()

add_normal('4. 维护建议', bold=True)
add_normal('（1）【建议1】')
add_normal('（2）【建议2】')
add_normal('（3）【建议3】')

doc.add_page_break()

# ============================================================
# 第四章：网络设备详细巡检
# ============================================================
add_heading_styled('三、网络设备巡检报告', level=2)

def add_device_inspection_table(device_type, device_name, ip, brand_model, serial, version,
                                 hardware_config, check_items, protocol_items, security_items):
    """添加设备巡检详情表"""
    # 设备标题
    add_heading_styled(f'{device_name}-{ip}', level=3)

    # 基本信息表
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'

    # 标题
    merge_cells(table, 0, 0, 0, 3)
    set_cell(table.rows[0].cells[0], f'{device_type}巡检报告', bold=True, size=Pt(12), align='center', bg='2E86C1')
    table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    merge_cells(table, 1, 0, 1, 3)
    set_cell(table.rows[1].cells[0], '基本信息检查', bold=True, bg='D6EAF8')

    set_cell(table.rows[2].cells[0], '主机型号', bold=True, bg='EBF5FB')
    set_cell(table.rows[2].cells[1], brand_model)
    set_cell(table.rows[2].cells[2], '序列号', bold=True, bg='EBF5FB')
    set_cell(table.rows[2].cells[3], serial)

    set_cell(table.rows[3].cells[0], '系统版本', bold=True, bg='EBF5FB')
    merge_cells(table, 3, 1, 3, 3)
    set_cell(table.rows[3].cells[1], version)

    set_cell(table.rows[4].cells[0], '硬件配置', bold=True, bg='EBF5FB')
    set_cell(table.rows[4].cells[1], '模块')
    merge_cells(table, 4, 2, 4, 3)
    set_cell(table.rows[4].cells[2], hardware_config.get('modules', ''))

    set_cell(table.rows[5].cells[0], '硬件配置', bold=True, bg='EBF5FB')
    set_cell(table.rows[5].cells[1], '接口数量')
    merge_cells(table, 5, 2, 5, 3)
    set_cell(table.rows[5].cells[2], hardware_config.get('interfaces', ''))

    set_cell(table.rows[6].cells[0], '硬件配置', bold=True, bg='EBF5FB')
    set_cell(table.rows[6].cells[1], '电源数量')
    merge_cells(table, 6, 2, 6, 3)
    set_cell(table.rows[6].cells[2], hardware_config.get('power', ''))

    set_cell(table.rows[7].cells[0], '硬件配置', bold=True, bg='EBF5FB')
    set_cell(table.rows[7].cells[1], '设备运行环境')
    merge_cells(table, 7, 2, 7, 3)
    set_cell(table.rows[7].cells[2], '√良好 □一般 □恶劣')

    doc.add_paragraph()

    # 设备运行状态检查
    table = doc.add_table(rows=len(check_items) + 2, cols=4)
    table.style = 'Table Grid'

    merge_cells(table, 0, 0, 0, 3)
    set_cell(table.rows[0].cells[0], '设备运行状态检查', bold=True, bg='2E86C1')
    table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    set_cell(table.rows[1].cells[0], '检查内容', bold=True, bg='D6EAF8', align='center')
    set_cell(table.rows[1].cells[1], '检查操作', bold=True, bg='D6EAF8', align='center')
    set_cell(table.rows[1].cells[2], '巡检结果', bold=True, bg='D6EAF8', align='center')
    set_cell(table.rows[1].cells[3], '结果说明', bold=True, bg='D6EAF8', align='center')

    for i, item in enumerate(check_items):
        set_cell(table.rows[i + 2].cells[0], item['name'], bold=True, size=FONT_SIZE_SMALL)
        set_cell(table.rows[i + 2].cells[1], item['command'], size=FONT_SIZE_SMALL)
        add_status_tag(table.rows[i + 2].cells[2], item.get('status', 'normal'))
        set_cell(table.rows[i + 2].cells[3], item.get('result', '【填写截图或日志】'), size=FONT_SIZE_SMALL)

    doc.add_paragraph()

    # 常见协议检查
    if protocol_items:
        table = doc.add_table(rows=len(protocol_items) + 2, cols=4)
        table.style = 'Table Grid'

        merge_cells(table, 0, 0, 0, 3)
        set_cell(table.rows[0].cells[0], '常见协议检查', bold=True, bg='2E86C1')
        table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        set_cell(table.rows[1].cells[0], '检查内容', bold=True, bg='D6EAF8', align='center')
        set_cell(table.rows[1].cells[1], '检查操作', bold=True, bg='D6EAF8', align='center')
        set_cell(table.rows[1].cells[2], '巡检结果', bold=True, bg='D6EAF8', align='center')
        set_cell(table.rows[1].cells[3], '结果说明', bold=True, bg='D6EAF8', align='center')

        for i, item in enumerate(protocol_items):
            set_cell(table.rows[i + 2].cells[0], item['name'], bold=True, size=FONT_SIZE_SMALL)
            set_cell(table.rows[i + 2].cells[1], item['command'], size=FONT_SIZE_SMALL)
            add_status_tag(table.rows[i + 2].cells[2], item.get('status', 'normal'))
            set_cell(table.rows[i + 2].cells[3], item.get('result', '【填写截图或日志】'), size=FONT_SIZE_SMALL)

        doc.add_paragraph()

    # 系统安全检查
    if security_items:
        table = doc.add_table(rows=len(security_items) + 2, cols=4)
        table.style = 'Table Grid'

        merge_cells(table, 0, 0, 0, 3)
        set_cell(table.rows[0].cells[0], '系统安全检查', bold=True, bg='2E86C1')
        table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        set_cell(table.rows[1].cells[0], '检查内容', bold=True, bg='D6EAF8', align='center')
        set_cell(table.rows[1].cells[1], '检查操作', bold=True, bg='D6EAF8', align='center')
        set_cell(table.rows[1].cells[2], '巡检结果', bold=True, bg='D6EAF8', align='center')
        set_cell(table.rows[1].cells[3], '结果说明', bold=True, bg='D6EAF8', align='center')

        for i, item in enumerate(security_items):
            set_cell(table.rows[i + 2].cells[0], item['name'], bold=True, size=FONT_SIZE_SMALL)
            set_cell(table.rows[i + 2].cells[1], item['command'], size=FONT_SIZE_SMALL)
            add_status_tag(table.rows[i + 2].cells[2], item.get('status', 'normal'))
            set_cell(table.rows[i + 2].cells[3], item.get('result', '【填写截图或日志】'), size=FONT_SIZE_SMALL)

    doc.add_page_break()

# ---- 路由器巡检示例 ----
add_device_inspection_table(
    device_type='路由器',
    device_name='生产边界路由器',
    ip='【IP地址】',
    brand_model='H3C SR6608-X',
    serial='【序列号】',
    version='Comware Software, Version 5.20.106',
    hardware_config={
        'modules': '主控板、业务板、接口板',
        'interfaces': '60千兆电口，12千兆光口',
        'power': '2（冗余）',
    },
    check_items=[
        {'name': '设备指示灯', 'command': '目测指示灯有无红灯', 'status': 'normal', 'result': '正常，无告警灯'},
        {'name': '设备运行时间', 'command': 'display version', 'status': 'normal', 'result': '运行XXX天，无异常重启'},
        {'name': 'CPU使用率', 'command': 'display cpu-usage', 'status': 'normal', 'result': '当前XX%，峰值XX%'},
        {'name': '内存使用率', 'command': 'display memory', 'status': 'normal', 'result': '使用率XX%'},
        {'name': '风扇状态', 'command': 'display fan', 'status': 'normal', 'result': '正常'},
        {'name': '电源状态', 'command': 'display power', 'status': 'normal', 'result': '正常'},
        {'name': '模块温度', 'command': 'display environment', 'status': 'normal', 'result': 'XX℃，正常'},
        {'name': '端口状态', 'command': 'display interface brief', 'status': 'normal', 'result': '所有业务端口UP'},
        {'name': '日志检查', 'command': 'display logbuffer', 'status': 'normal', 'result': '无异常报错'},
        {'name': '时钟检查', 'command': 'display clock', 'status': 'normal', 'result': '时钟正确'},
    ],
    protocol_items=[
        {'name': 'VRRP双机热备', 'command': 'display vrrp', 'status': 'normal', 'result': '主备状态正常'},
        {'name': 'OSPF路由协议', 'command': 'display ospf peer', 'status': 'normal', 'result': '邻居状态FULL'},
        {'name': 'BGP路由协议', 'command': 'display bgp peer', 'status': 'normal', 'result': '邻居状态Established'},
    ],
    security_items=[
        {'name': '远程登录', 'command': '检查登录方式、密码复杂度', 'status': 'normal', 'result': 'SSH登录，密码符合要求'},
        {'name': 'SNMP配置', 'command': 'display snmp-agent community', 'status': 'normal', 'result': '已修改默认团体字'},
        {'name': '日志服务器', 'command': 'display info-center', 'status': 'normal', 'result': '已配置日志服务器'},
        {'name': 'ACL访问控制', 'command': 'display acl all', 'status': 'normal', 'result': '已配置管理IP白名单'},
    ],
)

# ---- 交换机巡检示例 ----
add_device_inspection_table(
    device_type='交换机',
    device_name='内网核心交换机',
    ip='【IP地址】',
    brand_model='H3C S10508-V',
    serial='【序列号】',
    version='H3C Comware Software, Version 7.1',
    hardware_config={
        'modules': '主控板、业务板',
        'interfaces': '48千兆电口，24千兆光口，4万兆光口',
        'power': '2（冗余）',
    },
    check_items=[
        {'name': '设备指示灯', 'command': '目测指示灯有无红灯', 'status': 'normal', 'result': '正常'},
        {'name': '设备运行时间', 'command': 'display version', 'status': 'normal', 'result': '运行XXX天'},
        {'name': 'CPU使用率', 'command': 'display cpu-usage', 'status': 'normal', 'result': '当前XX%'},
        {'name': '内存使用率', 'command': 'display memory', 'status': 'normal', 'result': '使用率XX%'},
        {'name': '风扇状态', 'command': 'display fan', 'status': 'normal', 'result': '正常'},
        {'name': '电源状态', 'command': 'display power', 'status': 'normal', 'result': '正常'},
        {'name': '模块温度', 'command': 'display environment', 'status': 'normal', 'result': 'XX℃'},
        {'name': '端口状态', 'command': 'display interface brief', 'status': 'normal', 'result': '业务端口UP'},
        {'name': '日志检查', 'command': 'display logbuffer', 'status': 'normal', 'result': '无异常'},
        {'name': 'MAC地址表', 'command': 'display mac-address', 'status': 'normal', 'result': '正常'},
        {'name': 'ARP表', 'command': 'display arp', 'status': 'normal', 'result': '正常'},
    ],
    protocol_items=[
        {'name': 'VRRP双机热备', 'command': 'display vrrp', 'status': 'normal', 'result': '主备正常'},
        {'name': 'STP生成树', 'command': 'display stp', 'status': 'normal', 'result': '无环路'},
        {'name': 'OSPF路由', 'command': 'display ospf peer', 'status': 'normal', 'result': '邻居FULL'},
        {'name': '链路聚合', 'command': 'display link-aggregation', 'status': 'normal', 'result': '聚合正常'},
    ],
    security_items=[
        {'name': '远程登录', 'command': '检查登录方式', 'status': 'normal', 'result': 'SSH登录'},
        {'name': 'SNMP配置', 'command': 'display snmp-agent', 'status': 'normal', 'result': '已配置'},
        {'name': '端口安全', 'command': 'display port-security', 'status': 'normal', 'result': '已启用'},
    ],
)

# ---- 防火墙巡检示例 ----
add_device_inspection_table(
    device_type='防火墙',
    device_name='生产边界防火墙',
    ip='【IP地址】',
    brand_model='【品牌型号】',
    serial='【序列号】',
    version='【系统版本】',
    hardware_config={
        'modules': '主控板、业务板',
        'interfaces': '【接口数量】',
        'power': '2（冗余）',
    },
    check_items=[
        {'name': '设备指示灯', 'command': '目测指示灯', 'status': 'normal', 'result': '正常'},
        {'name': '设备运行时间', 'command': 'display version', 'status': 'normal', 'result': '运行XXX天'},
        {'name': 'CPU使用率', 'command': 'display cpu-usage', 'status': 'normal', 'result': '当前XX%'},
        {'name': '内存使用率', 'command': 'display memory', 'status': 'normal', 'result': '使用率XX%'},
        {'name': '会话数', 'command': 'display session statistics', 'status': 'normal', 'result': '当前XX万会话'},
        {'name': '风扇状态', 'command': 'display fan', 'status': 'normal', 'result': '正常'},
        {'name': '电源状态', 'command': 'display power', 'status': 'normal', 'result': '正常'},
        {'name': '模块温度', 'command': 'display environment', 'status': 'normal', 'result': 'XX℃'},
        {'name': '端口状态', 'command': 'display interface brief', 'status': 'normal', 'result': '正常'},
        {'name': '日志检查', 'command': 'display log', 'status': 'normal', 'result': '无异常'},
        {'name': '攻击日志', 'command': 'display attack-defense', 'status': 'normal', 'result': '无异常攻击'},
    ],
    protocol_items=[
        {'name': 'HA双机热备', 'command': 'display ha status', 'status': 'normal', 'result': '主备正常'},
        {'name': '路由表', 'command': 'display ip routing-table', 'status': 'normal', 'result': '路由正常'},
    ],
    security_items=[
        {'name': '安全策略', 'command': 'display security-policy', 'status': 'normal', 'result': '策略正常'},
        {'name': 'NAT策略', 'command': 'display nat policy', 'status': 'normal', 'result': 'NAT正常'},
        {'name': 'VPN隧道', 'command': 'display vpn tunnel', 'status': 'normal', 'result': '隧道正常'},
        {'name': '入侵防御', 'command': 'display ips status', 'status': 'normal', 'result': '已启用'},
    ],
)

# ============================================================
# 第五章：服务器详细巡检
# ============================================================
add_heading_styled('四、服务器巡检报告', level=2)

def add_server_inspection_table(server_type, server_name, ip, brand_model, serial, os_version,
                                 hardware_config, check_items):
    """添加服务器巡检详情表"""
    add_heading_styled(f'{server_name}-{ip}', level=3)

    # 基本信息表
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'

    merge_cells(table, 0, 0, 0, 3)
    set_cell(table.rows[0].cells[0], f'{server_type}巡检报告', bold=True, size=Pt(12), align='center', bg='2E86C1')
    table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    merge_cells(table, 1, 0, 1, 3)
    set_cell(table.rows[1].cells[0], '基本信息检查', bold=True, bg='D6EAF8')

    set_cell(table.rows[2].cells[0], '主机型号', bold=True, bg='EBF5FB')
    set_cell(table.rows[2].cells[1], brand_model)
    set_cell(table.rows[2].cells[2], '序列号', bold=True, bg='EBF5FB')
    set_cell(table.rows[2].cells[3], serial)

    set_cell(table.rows[3].cells[0], '操作系统', bold=True, bg='EBF5FB')
    merge_cells(table, 3, 1, 3, 3)
    set_cell(table.rows[3].cells[1], os_version)

    set_cell(table.rows[4].cells[0], '硬件配置', bold=True, bg='EBF5FB')
    set_cell(table.rows[4].cells[1], 'CPU')
    merge_cells(table, 4, 2, 4, 3)
    set_cell(table.rows[4].cells[2], hardware_config.get('cpu', ''))

    set_cell(table.rows[5].cells[0], '硬件配置', bold=True, bg='EBF5FB')
    set_cell(table.rows[5].cells[1], '内存')
    merge_cells(table, 5, 2, 5, 3)
    set_cell(table.rows[5].cells[2], hardware_config.get('memory', ''))

    set_cell(table.rows[6].cells[0], '硬件配置', bold=True, bg='EBF5FB')
    set_cell(table.rows[6].cells[1], '磁盘')
    merge_cells(table, 6, 2, 6, 3)
    set_cell(table.rows[6].cells[2], hardware_config.get('disk', ''))

    doc.add_paragraph()

    # 运行状态检查
    table = doc.add_table(rows=len(check_items) + 2, cols=4)
    table.style = 'Table Grid'

    merge_cells(table, 0, 0, 0, 3)
    set_cell(table.rows[0].cells[0], '设备运行状态检查', bold=True, bg='2E86C1')
    table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    set_cell(table.rows[1].cells[0], '检查内容', bold=True, bg='D6EAF8', align='center')
    set_cell(table.rows[1].cells[1], '检查操作', bold=True, bg='D6EAF8', align='center')
    set_cell(table.rows[1].cells[2], '巡检结果', bold=True, bg='D6EAF8', align='center')
    set_cell(table.rows[1].cells[3], '结果说明', bold=True, bg='D6EAF8', align='center')

    for i, item in enumerate(check_items):
        set_cell(table.rows[i + 2].cells[0], item['name'], bold=True, size=FONT_SIZE_SMALL)
        set_cell(table.rows[i + 2].cells[1], item['command'], size=FONT_SIZE_SMALL)
        add_status_tag(table.rows[i + 2].cells[2], item.get('status', 'normal'))
        set_cell(table.rows[i + 2].cells[3], item.get('result', '【填写截图或日志】'), size=FONT_SIZE_SMALL)

    doc.add_page_break()

# ---- 刀片服务器 ----
add_server_inspection_table(
    server_type='刀片服务器',
    server_name='HP UIS8000生产刀箱',
    ip='【IP地址】',
    brand_model='HP UIS8000',
    serial='【序列号】',
    os_version='VMware ESXi 7.0 / Windows Server 2019',
    hardware_config={
        'cpu': 'Intel Xeon E5-2680 v4 × 2',
        'memory': '256GB DDR4',
        'disk': '2TB SAS RAID1',
    },
    check_items=[
        {'name': '设备指示灯', 'command': '目测指示灯', 'status': 'normal', 'result': '正常'},
        {'name': '设备运行时间', 'command': 'uptime', 'status': 'normal', 'result': '运行XXX天'},
        {'name': 'CPU使用率', 'command': 'top/vmstat', 'status': 'normal', 'result': '当前XX%'},
        {'name': '内存使用率', 'command': 'free -m', 'status': 'normal', 'result': '使用率XX%'},
        {'name': '磁盘使用率', 'command': 'df -h', 'status': 'normal', 'result': '各分区<80%'},
        {'name': '系统负载', 'command': 'cat /proc/loadavg', 'status': 'normal', 'result': '负载XX'},
        {'name': '网络连接', 'command': 'ss -s', 'status': 'normal', 'result': '连接数正常'},
        {'name': '系统日志', 'command': 'journalctl -p err', 'status': 'normal', 'result': '无异常'},
        {'name': '风扇状态', 'command': 'iLO/iDRAC查看', 'status': 'normal', 'result': '正常'},
        {'name': '电源状态', 'command': 'iLO/iDRAC查看', 'status': 'normal', 'result': '正常'},
        {'name': '温度状态', 'command': 'iLO/iDRAC查看', 'status': 'normal', 'result': 'XX℃'},
    ],
)

# ---- 小型机 ----
add_server_inspection_table(
    server_type='小型机',
    server_name='IBM P750数据库节点',
    ip='【IP地址】',
    brand_model='IBM Power 750',
    serial='【序列号】',
    os_version='AIX 7.2 TL5 SP3',
    hardware_config={
        'cpu': 'POWER7+ 3.3GHz × 8',
        'memory': '256GB',
        'disk': 'SAS RAID10',
    },
    check_items=[
        {'name': '设备指示灯', 'command': '目测指示灯', 'status': 'normal', 'result': '正常'},
        {'name': '系统版本', 'command': 'oslevel -s', 'status': 'normal', 'result': 'AIX 7.2 TL5'},
        {'name': 'CPU使用率', 'command': 'vmstat 1 3', 'status': 'normal', 'result': '当前XX%'},
        {'name': '内存使用率', 'command': 'svmon -G', 'status': 'normal', 'result': '使用率XX%'},
        {'name': '文件系统', 'command': 'df -g', 'status': 'normal', 'result': '各FS<80%'},
        {'name': 'VG状态', 'command': 'lsvg -o', 'status': 'normal', 'result': 'VG正常'},
        {'name': 'Paging Space', 'command': 'lsps -a', 'status': 'normal', 'result': '使用率XX%'},
        {'name': 'HACMP状态', 'command': 'clstat', 'status': 'normal', 'result': '集群正常'},
        {'name': '网络接口', 'command': 'entstat -d ent0', 'status': 'normal', 'result': '正常'},
        {'name': '系统日志', 'command': 'errpt -a', 'status': 'normal', 'result': '无严重错误'},
        {'name': '硬件状态', 'command': 'lscfg -v', 'status': 'normal', 'result': '硬件正常'},
    ],
)

# ============================================================
# 第六章：存储设备详细巡检
# ============================================================
add_heading_styled('五、存储设备巡检报告', level=2)

def add_storage_inspection_table(storage_type, storage_name, ip, brand_model, serial, version,
                                  check_items):
    add_heading_styled(f'{storage_name}-{ip}', level=3)

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    merge_cells(table, 0, 0, 0, 3)
    set_cell(table.rows[0].cells[0], f'{storage_type}巡检报告', bold=True, size=Pt(12), align='center', bg='2E86C1')
    table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    merge_cells(table, 1, 0, 1, 3)
    set_cell(table.rows[1].cells[0], '基本信息检查', bold=True, bg='D6EAF8')

    set_cell(table.rows[2].cells[0], '设备型号', bold=True, bg='EBF5FB')
    set_cell(table.rows[2].cells[1], brand_model)
    set_cell(table.rows[2].cells[2], '序列号', bold=True, bg='EBF5FB')
    set_cell(table.rows[2].cells[3], serial)

    set_cell(table.rows[3].cells[0], '固件版本', bold=True, bg='EBF5FB')
    merge_cells(table, 3, 1, 3, 3)
    set_cell(table.rows[3].cells[1], version)

    set_cell(table.rows[4].cells[0], '设备运行环境', bold=True, bg='EBF5FB')
    merge_cells(table, 4, 1, 4, 3)
    set_cell(table.rows[4].cells[1], '√良好 □一般 □恶劣')

    doc.add_paragraph()

    table = doc.add_table(rows=len(check_items) + 2, cols=4)
    table.style = 'Table Grid'

    merge_cells(table, 0, 0, 0, 3)
    set_cell(table.rows[0].cells[0], '运行状态检查', bold=True, bg='2E86C1')
    table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    set_cell(table.rows[1].cells[0], '检查内容', bold=True, bg='D6EAF8', align='center')
    set_cell(table.rows[1].cells[1], '检查操作', bold=True, bg='D6EAF8', align='center')
    set_cell(table.rows[1].cells[2], '巡检结果', bold=True, bg='D6EAF8', align='center')
    set_cell(table.rows[1].cells[3], '结果说明', bold=True, bg='D6EAF8', align='center')

    for i, item in enumerate(check_items):
        set_cell(table.rows[i + 2].cells[0], item['name'], bold=True, size=FONT_SIZE_SMALL)
        set_cell(table.rows[i + 2].cells[1], item['command'], size=FONT_SIZE_SMALL)
        add_status_tag(table.rows[i + 2].cells[2], item.get('status', 'normal'))
        set_cell(table.rows[i + 2].cells[3], item.get('result', '【填写截图或日志】'), size=FONT_SIZE_SMALL)

    doc.add_page_break()

# ---- EMC存储 ----
add_storage_inspection_table(
    storage_type='磁盘阵列',
    storage_name='EMC VNX5600生产存储',
    ip='【IP地址】',
    brand_model='EMC VNX5600',
    serial='【序列号】',
    version='【固件版本】',
    check_items=[
        {'name': '系统状态', 'command': '查看存储管理界面', 'status': 'normal', 'result': '系统状态正常'},
        {'name': '控制器状态', 'command': '查看SPA/SPB状态', 'status': 'normal', 'result': '双控正常'},
        {'name': '存储池使用率', 'command': '查看Pool使用率', 'status': 'normal', 'result': '使用率XX%'},
        {'name': 'LUN状态', 'command': '查看LUN状态', 'status': 'normal', 'result': '所有LUN正常'},
        {'name': '硬盘状态', 'command': '查看磁盘状态', 'status': 'normal', 'result': '无故障盘'},
        {'name': '电源状态', 'command': '查看电源模块', 'status': 'normal', 'result': '正常'},
        {'name': '风扇状态', 'command': '查看风扇模块', 'status': 'normal', 'result': '正常'},
        {'name': '缓存状态', 'command': '查看缓存状态', 'status': 'normal', 'result': '缓存正常'},
        {'name': '告警信息', 'command': '查看系统告警', 'status': 'normal', 'result': '无严重告警'},
    ],
)

# ---- 光纤交换机 ----
add_storage_inspection_table(
    storage_type='光纤交换机',
    storage_name='EMC 6505B生产光交',
    ip='【IP地址】',
    brand_model='EMC DS6505B',
    serial='【序列号】',
    version='【固件版本】',
    check_items=[
        {'name': '交换机状态', 'command': 'switchshow', 'status': 'normal', 'result': '状态正常'},
        {'name': '端口状态', 'command': 'portshow', 'status': 'normal', 'result': '所有端口Online'},
        {'name': 'SFP信息', 'command': 'sfpshow all', 'status': 'normal', 'result': 'SFP正常'},
        {'name': '错误统计', 'command': 'porterrshow', 'status': 'normal', 'result': '无CRC错误'},
        {'name': '固件版本', 'command': 'firmwareshow', 'status': 'normal', 'result': '版本一致'},
        {'name': '温度状态', 'command': 'tempshow', 'status': 'normal', 'result': 'XX℃'},
        {'name': '日志检查', 'command': 'errdump', 'status': 'normal', 'result': '无异常'},
    ],
)

# ============================================================
# 第七章：数据库巡检
# ============================================================
add_heading_styled('六、数据库系统巡检报告', level=2)
add_normal('本次数据库常规检查的数据收集主要集中在【YYYY年MM月】，我们尽可能把重要的信息收集起来进行分析。')
add_normal('以下是本次检查的总结：')

add_normal('1. 主要数据库问题汇总', bold=True)

table = doc.add_table(rows=6, cols=5)
table.style = 'Table Grid'
headers = ['序号', '数据库名称', 'IP地址', '问题描述', '严重程度']
for i, h in enumerate(headers):
    set_cell(table.rows[0].cells[i], h, bold=True, bg='E74C3C', align='center', size=FONT_SIZE_SMALL)
    table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
for r in range(1, 6):
    set_cell(table.rows[r].cells[0], str(r), align='center', size=FONT_SIZE_SMALL)
    for c in range(1, 5):
        set_cell(table.rows[r].cells[c], '【填写】', size=FONT_SIZE_SMALL)

doc.add_paragraph()

add_normal('2. 巡检信息汇总', bold=True)

table = doc.add_table(rows=5, cols=6)
table.style = 'Table Grid'
headers = ['数据库名称', '版本', 'IP地址', '实例状态', '表空间使用率', '健康度']
for i, h in enumerate(headers):
    set_cell(table.rows[0].cells[i], h, bold=True, bg='2E86C1', align='center', size=FONT_SIZE_SMALL)
    table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
for r in range(1, 5):
    for c in range(6):
        set_cell(table.rows[r].cells[c], '【填写】', size=FONT_SIZE_SMALL, align='center')

doc.add_page_break()

# ============================================================
# 第八章：安全设备巡检
# ============================================================
add_heading_styled('七、安全设备巡检报告', level=2)

add_device_inspection_table(
    device_type='WAF防火墙',
    device_name='外网WAF防火墙',
    ip='【IP地址】',
    brand_model='【品牌型号】',
    serial='【序列号】',
    version='【系统版本】',
    hardware_config={
        'modules': '主控板、业务板',
        'interfaces': '【接口数量】',
        'power': '2（冗余）',
    },
    check_items=[
        {'name': '设备指示灯', 'command': '目测指示灯', 'status': 'normal', 'result': '正常'},
        {'name': '设备运行时间', 'command': 'display version', 'status': 'normal', 'result': '运行XXX天'},
        {'name': 'CPU使用率', 'command': 'display cpu-usage', 'status': 'normal', 'result': '当前XX%'},
        {'name': '内存使用率', 'command': 'display memory', 'status': 'normal', 'result': '使用率XX%'},
        {'name': 'WAF策略状态', 'command': '查看WAF策略', 'status': 'normal', 'result': '策略正常'},
        {'name': '攻击日志', 'command': '查看攻击日志', 'status': 'normal', 'result': '无异常攻击'},
        {'name': '证书状态', 'command': '查看SSL证书', 'status': 'normal', 'result': '证书有效'},
        {'name': '日志检查', 'command': 'display log', 'status': 'normal', 'result': '无异常'},
    ],
    protocol_items=[],
    security_items=[
        {'name': '安全策略', 'command': '查看安全策略', 'status': 'normal', 'result': '策略正常'},
        {'name': '访问控制', 'command': '查看ACL', 'status': 'normal', 'result': '已配置'},
    ],
)

# ============================================================
# 第九章：虚拟化巡检
# ============================================================
add_heading_styled('八、虚拟化平台巡检报告', level=2)

add_heading_styled('VMware生产集群-【IP】', level=3)

table = doc.add_table(rows=12, cols=4)
table.style = 'Table Grid'

merge_cells(table, 0, 0, 0, 3)
set_cell(table.rows[0].cells[0], 'VMware虚拟化巡检报告', bold=True, size=Pt(12), align='center', bg='2E86C1')
table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

merge_cells(table, 1, 0, 1, 3)
set_cell(table.rows[1].cells[0], '基本信息检查', bold=True, bg='D6EAF8')

set_cell(table.rows[2].cells[0], 'vCenter版本', bold=True, bg='EBF5FB')
merge_cells(table, 2, 1, 2, 3)
set_cell(table.rows[2].cells[1], '【版本号】')

set_cell(table.rows[3].cells[0], 'ESXi版本', bold=True, bg='EBF5FB')
merge_cells(table, 3, 1, 3, 3)
set_cell(table.rows[3].cells[1], '【版本号】')

set_cell(table.rows[4].cells[0], '集群名称', bold=True, bg='EBF5FB')
merge_cells(table, 4, 1, 4, 3)
set_cell(table.rows[4].cells[1], '【集群名称】')

set_cell(table.rows[5].cells[0], '主机数量', bold=True, bg='EBF5FB')
set_cell(table.rows[5].cells[1], '【N】台')
set_cell(table.rows[5].cells[2], '虚拟机数量', bold=True, bg='EBF5FB')
set_cell(table.rows[5].cells[3], '【N】台')

merge_cells(table, 6, 0, 6, 3)
set_cell(table.rows[6].cells[0], '运行状态检查', bold=True, bg='D6EAF8')

checks = [
    ['vCenter服务状态', '检查vCenter服务', '正常', '所有服务运行正常'],
    ['ESXi主机状态', '检查所有ESXi主机', '正常', '所有主机已连接'],
    ['集群资源使用', 'CPU/内存/存储使用率', '正常', 'CPU XX%, 内存 XX%'],
    ['数据存储状态', '检查datastore使用率', '正常', '各datastore<80%'],
    ['HA/DRS状态', '检查HA和DRS配置', '正常', 'HA/DRS已启用'],
]

for i, check in enumerate(checks):
    set_cell(table.rows[7 + i].cells[0], check[0], bold=True, size=FONT_SIZE_SMALL)
    set_cell(table.rows[7 + i].cells[1], check[1], size=FONT_SIZE_SMALL)
    add_status_tag(table.rows[7 + i].cells[2], 'normal')
    set_cell(table.rows[7 + i].cells[3], check[3], size=FONT_SIZE_SMALL)

doc.add_page_break()

# ============================================================
# 第十章：备份系统巡检
# ============================================================
add_heading_styled('九、备份系统巡检报告', level=2)

add_heading_styled('NetBackup备份服务器-【IP】', level=3)

table = doc.add_table(rows=10, cols=4)
table.style = 'Table Grid'

merge_cells(table, 0, 0, 0, 3)
set_cell(table.rows[0].cells[0], 'NetBackup巡检报告', bold=True, size=Pt(12), align='center', bg='2E86C1')
table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

merge_cells(table, 1, 0, 1, 3)
set_cell(table.rows[1].cells[0], '基本信息检查', bold=True, bg='D6EAF8')

set_cell(table.rows[2].cells[0], '备份服务器', bold=True, bg='EBF5FB')
set_cell(table.rows[2].cells[1], '【服务器名称】')
set_cell(table.rows[2].cells[2], 'IP地址', bold=True, bg='EBF5FB')
set_cell(table.rows[2].cells[3], '【IP】')

set_cell(table.rows[3].cells[0], 'NBU版本', bold=True, bg='EBF5FB')
merge_cells(table, 3, 1, 3, 3)
set_cell(table.rows[3].cells[1], '【版本号】')

merge_cells(table, 4, 0, 4, 3)
set_cell(table.rows[4].cells[0], '运行状态检查', bold=True, bg='D6EAF8')

backup_checks = [
    ['备份服务状态', '检查nbu服务', '正常', '所有服务运行正常'],
    ['近期备份任务', '查看备份任务日志', '正常', '所有任务成功完成'],
    ['磁带库状态', '检查磁带库', '正常', '磁带库正常'],
    ['存储单元状态', '检查STU', '正常', '所有STU在线'],
    ['备份策略', '检查备份策略配置', '正常', '策略配置正确'],
]

for i, check in enumerate(backup_checks):
    set_cell(table.rows[5 + i].cells[0], check[0], bold=True, size=FONT_SIZE_SMALL)
    set_cell(table.rows[5 + i].cells[1], check[1], size=FONT_SIZE_SMALL)
    add_status_tag(table.rows[5 + i].cells[2], 'normal')
    set_cell(table.rows[5 + i].cells[3], check[3], size=FONT_SIZE_SMALL)

doc.add_page_break()

# ============================================================
# 附录：巡检工具和命令参考
# ============================================================
add_heading_styled('附录：巡检工具和命令参考', level=1)

add_normal('1. 网络设备巡检命令', bold=True)
table = doc.add_table(rows=11, cols=3)
table.style = 'Table Grid'
set_cell(table.rows[0].cells[0], '检查项', bold=True, bg='2E86C1', align='center')
set_cell(table.rows[0].cells[1], 'H3C/华为命令', bold=True, bg='2E86C1', align='center')
set_cell(table.rows[0].cells[2], '思科命令', bold=True, bg='2E86C1', align='center')
table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
table.rows[0].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
table.rows[0].cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

cmds = [
    ['版本信息', 'display version', 'show version'],
    ['CPU使用率', 'display cpu-usage', 'show processes cpu'],
    ['内存使用率', 'display memory', 'show memory statistics'],
    ['接口状态', 'display interface brief', 'show ip interface brief'],
    ['路由表', 'display ip routing-table', 'show ip route'],
    ['风扇状态', 'display fan', 'show environment fan'],
    ['电源状态', 'display power', 'show environment power'],
    ['温度信息', 'display environment', 'show environment temperature'],
    ['日志信息', 'display logbuffer', 'show logging'],
    ['时钟信息', 'display clock', 'show clock'],
]
for i, cmd in enumerate(cmds):
    set_cell(table.rows[i+1].cells[0], cmd[0], bold=True, size=FONT_SIZE_SMALL)
    set_cell(table.rows[i+1].cells[1], cmd[1], size=FONT_SIZE_SMALL)
    set_cell(table.rows[i+1].cells[2], cmd[2], size=FONT_SIZE_SMALL)

doc.add_paragraph()

add_normal('2. Linux服务器巡检命令', bold=True)
table = doc.add_table(rows=11, cols=2)
table.style = 'Table Grid'
set_cell(table.rows[0].cells[0], '检查项', bold=True, bg='2E86C1', align='center')
set_cell(table.rows[0].cells[1], '命令', bold=True, bg='2E86C1', align='center')
table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
table.rows[0].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

linux_cmds = [
    ['系统版本', 'cat /etc/os-release && uname -a'],
    ['运行时间', 'uptime'],
    ['CPU使用率', 'top -bn1 | grep Cpu'],
    ['内存使用率', 'free -m'],
    ['磁盘使用率', 'df -h'],
    ['系统负载', 'cat /proc/loadavg'],
    ['网络连接', 'ss -s'],
    ['进程数', 'ps aux | wc -l'],
    ['系统日志', 'journalctl -p err --since "24 hours ago"'],
    ['NTP同步', 'timedatectl status'],
]
for i, cmd in enumerate(linux_cmds):
    set_cell(table.rows[i+1].cells[0], cmd[0], bold=True, size=FONT_SIZE_SMALL)
    set_cell(table.rows[i+1].cells[1], cmd[1], size=FONT_SIZE_SMALL)

# ============================================================
# 设置页面边距
# ============================================================
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ============================================================
# 保存
# ============================================================
output_path = '/home/work/.openclaw/workspace/acdante-itops-inspection/templates/综合巡检报告模板.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'模板已生成: {output_path}')
print(f'文件大小: {os.path.getsize(output_path)} bytes')
