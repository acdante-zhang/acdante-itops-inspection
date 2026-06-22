"""
Acdante ITOps - DOCX 报告生成器
基于 python-docx 生成贴近实际巡检模板的 Word 报告
"""

import os
import io
from datetime import datetime
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from . import (
    InspectionReportData, ReportConfig, OUTPUT_DIR,
    DeviceInspectionResult, DeviceInfo, InspectionSection, InspectionCheckItem,
)
from .device_templates import (
    get_template_for_device, SectionDef, CheckItemDef,
    DEVICE_TYPE_TEMPLATE_MAP,
)


class DocxGenerator:
    """DOCX报告生成器 - 模板化版本"""

    # 字体
    FONT_NAME = '微软雅黑'
    FONT_NAME_MONO = 'Consolas'

    # 颜色
    COLOR_PRIMARY = RGBColor(0x0E, 0x74, 0x9A)
    COLOR_DARK = RGBColor(0x1E, 0x29, 0x3B)
    COLOR_OK = RGBColor(0x16, 0xA3, 0x4A)
    COLOR_WARNING = RGBColor(0xE5, 0xA7, 0x00)
    COLOR_CRITICAL = RGBColor(0xDC, 0x26, 0x26)
    COLOR_TEXT = RGBColor(0x33, 0x33, 0x33)
    COLOR_MUTED = RGBColor(0x6B, 0x72, 0x80)
    COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

    def __init__(self, data: InspectionReportData):
        self.data = data
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """设置文档全局样式"""
        style = self.doc.styles['Normal']
        style.font.name = self.FONT_NAME
        style.font.size = Pt(10)
        style.font.color.rgb = self.COLOR_TEXT
        style.element.rPr.rFonts.set(qn('w:eastAsia'), self.FONT_NAME)

        for section in self.doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

    # ============================================================
    # 通用工具方法
    # ============================================================

    def _set_run_font(self, run, name=None, size=None, bold=None, color=None):
        """统一设置 run 字体"""
        run.font.name = name or self.FONT_NAME
        run.element.rPr.rFonts.set(qn('w:eastAsia'), name or self.FONT_NAME)
        if size:
            run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold
        if color:
            run.font.color.rgb = color

    def _add_paragraph(self, text: str, bold=False, size=10, color=None,
                       alignment=None, space_before=None, space_after=None):
        """添加段落"""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        self._set_run_font(run, size=size, bold=bold, color=color)
        if alignment is not None:
            para.alignment = alignment
        if space_before is not None:
            para.paragraph_format.space_before = Pt(space_before)
        if space_after is not None:
            para.paragraph_format.space_after = Pt(space_after)
        return para

    def _add_heading(self, text: str, level: int = 1):
        """添加标题"""
        heading = self.doc.add_heading(text, level=level)
        for run in heading.runs:
            self._set_run_font(run)
        return heading

    def _set_cell_shading(self, cell, color_hex: str):
        """设置单元格背景色"""
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def _set_cell_text(self, cell, text: str, bold=False, size=9, color=None,
                       alignment=None, font_name=None):
        """设置单元格文字"""
        cell.text = ''
        para = cell.paragraphs[0]
        run = para.add_run(str(text) if text else '')
        self._set_run_font(run, name=font_name, size=size, bold=bold, color=color)
        if alignment is not None:
            para.alignment = alignment

    def _merge_cells_in_row(self, table, row_idx, col_start, col_end):
        """合并行中的单元格"""
        cell_start = table.cell(row_idx, col_start)
        cell_end = table.cell(row_idx, col_end)
        return cell_start.merge(cell_end)

    # ============================================================
    # 封面页生成
    # ============================================================

    def _generate_cover(self):
        """生成封面页 - 匹配模板样式"""
        config = self.data.config

        # 空行调整位置
        for _ in range(4):
            self.doc.add_paragraph()

        # 客户单位名称
        if config.company_name:
            self._add_paragraph(
                config.company_name,
                bold=True, size=24,
                alignment=WD_ALIGN_PARAGRAPH.CENTER
            )

        # 报告标题
        self._add_paragraph(
            config.title or "巡检服务报告",
            bold=True, size=24,
            alignment=WD_ALIGN_PARAGRAPH.CENTER
        )

        # 副标题
        if config.subtitle:
            self._add_paragraph(
                config.subtitle,
                bold=True, size=14,
                alignment=WD_ALIGN_PARAGRAPH.CENTER
            )

        # 巡检周期 (大字)
        if config.inspection_period:
            self._add_paragraph(
                config.inspection_period,
                bold=True, size=18,
                alignment=WD_ALIGN_PARAGRAPH.CENTER
            )

        # 空行
        for _ in range(2):
            self.doc.add_paragraph()

        # 服务公司
        if config.service_team:
            self._add_paragraph(
                config.service_team,
                bold=True, size=14,
                alignment=WD_ALIGN_PARAGRAPH.CENTER
            )

        # 分页
        self.doc.add_page_break()

    # ============================================================
    # 文档控制页
    # ============================================================

    def _generate_doc_control(self):
        """生成文档控制页"""
        self._add_heading("文档控制", level=2)

        # 客户信息表
        config = self.data.config
        self._add_paragraph("版本修改记录", bold=True, size=14, space_after=6)

        # 保密声明
        if config.company_name:
            self._add_paragraph(
                f"此文档仅供{config.company_name}相关人员审阅",
                size=10, color=self.COLOR_MUTED
            )
            self._add_paragraph(
                "不得向与此无关的个人或机构传阅或复制",
                size=10, color=self.COLOR_MUTED, space_after=12
            )

        # 客户联系信息表
        info_table = self.doc.add_table(rows=3, cols=4)
        info_table.style = 'Table Grid'
        info_data = [
            ["客户联系人", config.client_contact or "", "客户电话", config.client_phone or ""],
            ["巡检报告编写", "/".join(config.report_authors) if config.report_authors else "",
             "编写日期", config.author_date or ""],
            ["巡检报告审批", config.report_reviewer or "", "审批日期", config.reviewer_date or ""],
        ]
        for r, row_data in enumerate(info_data):
            for c, val in enumerate(row_data):
                bold = c in (0, 2)
                self._set_cell_text(info_table.cell(r, c), val, bold=bold, size=9)

        self.doc.add_paragraph()

        # 修改记录表
        self._add_paragraph("修改记录", bold=True, size=11, space_after=6)
        rev_table = self.doc.add_table(rows=2, cols=4)
        rev_table.style = 'Table Grid'
        rev_headers = ["日期", "作者", "版本", "修改记录"]
        for c, h in enumerate(rev_headers):
            self._set_cell_text(rev_table.cell(0, c), h, bold=True, size=9,
                                alignment=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell_shading(rev_table.cell(0, c), "F1F5F9")
        rev_data = [
            config.author_date or datetime.now().strftime("%Y.%m.%d"),
            "/".join(config.report_authors) if config.report_authors else "",
            "V1.0", "创建文档"
        ]
        for c, val in enumerate(rev_data):
            self._set_cell_text(rev_table.cell(1, c), val, size=9)

        self.doc.add_paragraph()

        # 审阅记录表
        self._add_paragraph("审阅记录", bold=True, size=11, space_after=6)
        review_table = self.doc.add_table(rows=2, cols=2)
        review_table.style = 'Table Grid'
        self._set_cell_text(review_table.cell(0, 0), "审阅人", bold=True, size=9,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._set_cell_text(review_table.cell(0, 1), "备注", bold=True, size=9,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._set_cell_shading(review_table.cell(0, 0), "F1F5F9")
        self._set_cell_shading(review_table.cell(0, 1), "F1F5F9")
        self._set_cell_text(review_table.cell(1, 0), config.report_reviewer or "", size=9)
        self._set_cell_text(review_table.cell(1, 1), "服务质量监督", size=9)

        self.doc.add_paragraph()

        # 分发记录表
        self._add_paragraph("分发记录", bold=True, size=11, space_after=6)
        dist_table = self.doc.add_table(rows=2, cols=3)
        dist_table.style = 'Table Grid'
        dist_headers = ["拷贝No.", "姓名", "单位"]
        for c, h in enumerate(dist_headers):
            self._set_cell_text(dist_table.cell(0, c), h, bold=True, size=9,
                                alignment=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell_shading(dist_table.cell(0, c), "F1F5F9")
        self._set_cell_text(dist_table.cell(1, 0), "1", size=9)
        self._set_cell_text(dist_table.cell(1, 1), config.client_contact or "", size=9)
        self._set_cell_text(dist_table.cell(1, 2), config.company_name or "", size=9)

        self.doc.add_page_break()

    # ============================================================
    # 目录页
    # ============================================================

    def _generate_toc(self):
        """生成目录页"""
        self._add_heading("目录", level=1)
        # 添加 TOC 域代码 (需要在 Word 中按 F9 刷新)
        para = self.doc.add_paragraph()
        run = para.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        run2 = para.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
        run2._r.append(instrText)
        run3 = para.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run3._r.append(fldChar2)
        run4 = para.add_run('（请在Word中右键点击此处，选择"更新域"以生成目录）')
        self._set_run_font(run4, size=9, color=self.COLOR_MUTED)
        run5 = para.add_run()
        fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run5._r.append(fldChar3)

        self.doc.add_page_break()

    # ============================================================
    # 巡检任务信息
    # ============================================================

    def _generate_task_info(self):
        """生成巡检任务信息总表"""
        self._add_heading("巡检服务报告", level=1)

        config = self.data.config

        # 巡检任务信息表 (大表)
        task_table = self.doc.add_table(rows=11, cols=15)
        task_table.style = 'Table Grid'

        # Row 0: 标题行 - 合并所有单元格
        self._merge_cells_in_row(task_table, 0, 0, 14)
        self._set_cell_text(task_table.cell(0, 0), "巡检任务信息", bold=True, size=11,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._set_cell_shading(task_table.cell(0, 0), "F1F5F9")

        # Row 1: 用户单位
        self._merge_cells_in_row(task_table, 1, 0, 1)
        self._set_cell_text(task_table.cell(1, 0), "用户单位", bold=True, size=9)
        self._merge_cells_in_row(task_table, 1, 2, 14)
        self._set_cell_text(task_table.cell(1, 2), config.company_name or "", size=9)

        # Row 2: 巡检时间 + 巡检周期
        self._merge_cells_in_row(task_table, 2, 0, 1)
        self._set_cell_text(task_table.cell(2, 0), "巡检时间", bold=True, size=9)
        self._merge_cells_in_row(task_table, 2, 2, 5)
        self._set_cell_text(task_table.cell(2, 2), config.inspection_date_range or "", size=9)
        self._merge_cells_in_row(task_table, 2, 6, 10)
        self._set_cell_text(task_table.cell(2, 6), "巡检周期", bold=True, size=9)
        self._merge_cells_in_row(task_table, 2, 11, 14)
        self._set_cell_text(task_table.cell(2, 11), config.inspection_period or "", size=9)

        # Row 3: 技术服务团队
        self._merge_cells_in_row(task_table, 3, 0, 14)
        self._set_cell_text(task_table.cell(3, 0), config.service_team or "", bold=True, size=9,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Row 4: 巡检工程师表头
        eng_headers = [
            (0, 1, "巡检工程师"), (2, 5, "职责"), (6, 9, "姓名"), (10, 14, "联系方式")
        ]
        for cs, ce, text in eng_headers:
            self._merge_cells_in_row(task_table, 4, cs, ce)
            self._set_cell_text(task_table.cell(4, cs), text, bold=True, size=9,
                                alignment=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell_shading(task_table.cell(4, cs), "F1F5F9")

        # Row 5-6: 工程师数据
        engineers = config.engineers or [{"role": "", "name": "", "phone": ""}]
        for idx, eng in enumerate(engineers[:2]):
            row = 5 + idx
            self._merge_cells_in_row(task_table, row, 0, 1)
            self._set_cell_text(task_table.cell(row, 0), "巡检工程师", bold=True, size=9)
            self._merge_cells_in_row(task_table, row, 2, 5)
            self._set_cell_text(task_table.cell(row, 2), eng.get("role", ""), size=9)
            self._merge_cells_in_row(task_table, row, 6, 9)
            self._set_cell_text(task_table.cell(row, 6), eng.get("name", ""), size=9)
            self._merge_cells_in_row(task_table, row, 10, 14)
            self._set_cell_text(task_table.cell(row, 10), eng.get("phone", ""), size=9)

        # Row 7: 巡检服务内容标题
        self._merge_cells_in_row(task_table, 7, 0, 14)
        self._set_cell_text(task_table.cell(7, 0), "巡检服务内容", bold=True, size=9,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._set_cell_shading(task_table.cell(7, 0), "F1F5F9")

        # Row 8: 服务内容
        self._merge_cells_in_row(task_table, 8, 0, 14)
        self._set_cell_text(task_table.cell(8, 0),
                            config.service_content or "根据维保合同设备清单，按照月度巡检计划到达客户现场，进行网络、主机、存储和数据库的巡检和优化调整",
                            size=9)

        # Row 9: 巡检结果标题
        self._merge_cells_in_row(task_table, 9, 0, 14)
        self._set_cell_text(task_table.cell(9, 0), "巡检结果及维护建议", bold=True, size=9,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._set_cell_shading(task_table.cell(9, 0), "F1F5F9")

        # Row 10: 巡检结果摘要
        self._merge_cells_in_row(task_table, 10, 0, 14)
        self._set_cell_text(task_table.cell(10, 0), self.data.summary or "", size=9)

        self.doc.add_paragraph()

        # 硬件系统巡检清单表
        self._add_paragraph("一、硬件系统巡检清单：问题和概要见上表总结", bold=True, size=10)
        self._generate_device_inventory_table()

        self.doc.add_page_break()

    def _generate_device_inventory_table(self):
        """生成设备清单表"""
        devices = [dr.device for dr in self.data.device_results] if self.data.device_results else []
        if not devices:
            return

        headers = ["序号", "类型", "设备类型", "品牌型号", "数量", "序列号", "IP", "备注"]
        table = self.doc.add_table(rows=1 + len(devices), cols=len(headers))
        table.style = 'Table Grid'

        # 表头
        for c, h in enumerate(headers):
            self._set_cell_text(table.cell(0, c), h, bold=True, size=8,
                                alignment=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell_shading(table.cell(0, c), "F1F5F9")

        # 数据行
        for r, dev in enumerate(devices):
            row = r + 1
            self._set_cell_text(table.cell(row, 0), str(r + 1), size=8)
            self._set_cell_text(table.cell(row, 1), dev.device_type.value, size=8)
            self._set_cell_text(table.cell(row, 2), dev.name, size=8)
            self._set_cell_text(table.cell(row, 3), f"{dev.brand} {dev.model}".strip(), size=8)
            self._set_cell_text(table.cell(row, 4), "1", size=8)
            self._set_cell_text(table.cell(row, 5), dev.serial_number or "N/A", size=8)
            self._set_cell_text(table.cell(row, 6), dev.ip or "N/A", size=8)
            self._set_cell_text(table.cell(row, 7), dev.role or "", size=8)

    # ============================================================
    # 设备巡检表格生成
    # ============================================================

    def _generate_device_section(self, device_result: DeviceInspectionResult):
        """生成单台设备的巡检报告段落"""
        device = device_result.device
        template = get_template_for_device(device.device_type.value)

        # 标题 (H2)
        title = f"{device.name}-{device.ip}" if device.ip else device.name
        self._add_heading(title, level=2)

        # 获取模板定义
        if template:
            self._generate_device_table_from_template(device_result, template)
        else:
            self._generate_device_table_generic(device_result)

    def _generate_device_table_from_template(self, device_result: DeviceInspectionResult,
                                              template: dict):
        """根据模板定义生成设备巡检表"""
        device = device_result.device
        sections = device_result.sections
        template_sections = template.get("sections", [])

        # 计算总行数
        total_rows = 1  # 标题行
        for sec_def in template_sections:
            if sec_def.section_type == "info":
                total_rows += 1  # 分区标题行
                total_rows += len(sec_def.info_fields)
            elif sec_def.section_type == "check":
                total_rows += 1  # 分区标题行
                total_rows += 1  # 表头行
                total_rows += len(sec_def.check_items)
            elif sec_def.section_type == "check_grid":
                total_rows += 1  # 分区标题行
                total_rows += 1  # 表头行
                total_rows += (len(sec_def.check_items) + 1) // 2  # 2列布局
            elif sec_def.section_type == "protocol":
                total_rows += 1  # 分区标题行
                total_rows += 1  # 表头行
                total_rows += len(sec_def.check_items)
            elif sec_def.section_type == "security":
                total_rows += 1  # 分区标题行
                total_rows += 1  # 表头行
                total_rows += len(sec_def.check_items)
            elif sec_def.section_type == "blade_slots":
                total_rows += 1  # 分区标题行
                total_rows += 1  # 表头行
                total_rows += 4  # 默认4个刀片槽位

        # 创建表格 (4列)
        table = self.doc.add_table(rows=total_rows, cols=4)
        table.style = 'Table Grid'

        # 构建检查项结果映射
        result_map = {}
        for sec in sections:
            for item in sec.items:
                result_map[item.name] = item

        row_idx = 0

        # Row 0: 标题行 - 合并
        self._merge_cells_in_row(table, 0, 0, 3)
        self._set_cell_text(table.cell(0, 0), template["title"], bold=True, size=11,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._set_cell_shading(table.cell(0, 0), "E8F4F8")
        row_idx = 1

        for sec_def in template_sections:
            if sec_def.section_type == "info":
                # 基本信息检查分区
                self._merge_cells_in_row(table, row_idx, 0, 3)
                self._set_cell_text(table.cell(row_idx, 0), sec_def.title, bold=True, size=10)
                self._set_cell_shading(table.cell(row_idx, 0), "F1F5F9")
                row_idx += 1

                # 填充基本信息
                info_values = self._get_info_values(device, sec_def)
                for field_label, field_placeholder in sec_def.info_fields:
                    if field_label:
                        self._set_cell_text(table.cell(row_idx, 0), field_label, bold=True, size=9)
                        # 查找对应的设备信息值
                        val = info_values.get(field_label, "")
                        self._merge_cells_in_row(table, row_idx, 1, 2)
                        self._set_cell_text(table.cell(row_idx, 1), val, size=9)
                        self._set_cell_text(table.cell(row_idx, 3), "", size=9)
                    else:
                        # 子项 (如接口数量、电源数量)
                        self._set_cell_text(table.cell(row_idx, 0), "", size=9)
                        self._set_cell_text(table.cell(row_idx, 1), field_placeholder, size=9)
                        val = info_values.get(field_placeholder, "")
                        self._merge_cells_in_row(table, row_idx, 2, 3)
                        self._set_cell_text(table.cell(row_idx, 2), val, size=9)
                    row_idx += 1

            elif sec_def.section_type in ("check", "protocol", "security"):
                # 检查项分区
                self._merge_cells_in_row(table, row_idx, 0, 3)
                self._set_cell_text(table.cell(row_idx, 0), sec_def.title, bold=True, size=10)
                self._set_cell_shading(table.cell(row_idx, 0), "F1F5F9")
                row_idx += 1

                # 表头
                self._set_cell_text(table.cell(row_idx, 0), "检查内容", bold=True, size=9,
                                    alignment=WD_ALIGN_PARAGRAPH.CENTER)
                self._set_cell_text(table.cell(row_idx, 1), "检查操作", bold=True, size=9,
                                    alignment=WD_ALIGN_PARAGRAPH.CENTER)
                self._merge_cells_in_row(table, row_idx, 2, 3)
                self._set_cell_text(table.cell(row_idx, 2), "巡检结果", bold=True, size=9,
                                    alignment=WD_ALIGN_PARAGRAPH.CENTER)
                self._set_cell_shading(table.cell(row_idx, 0), "F8FAFC")
                self._set_cell_shading(table.cell(row_idx, 1), "F8FAFC")
                self._set_cell_shading(table.cell(row_idx, 2), "F8FAFC")
                row_idx += 1

                # 检查项
                for check_def in sec_def.check_items:
                    self._set_cell_text(table.cell(row_idx, 0), check_def.name, size=9)
                    self._set_cell_text(table.cell(row_idx, 1), check_def.operation, size=8)

                    # 查找实际结果
                    actual = result_map.get(check_def.name)
                    if actual:
                        result_text = f"结果说明: {actual.raw_value or actual.result}"
                        status_text = self._get_status_label(actual.status)
                        result_display = f"{status_text} {result_text}"
                    else:
                        result_display = check_def.result_hint

                    self._merge_cells_in_row(table, row_idx, 2, 3)
                    self._set_cell_text(table.cell(row_idx, 2), result_display, size=9)
                    row_idx += 1

            elif sec_def.section_type == "check_grid":
                # 2列网格巡检 (用于存储设备)
                self._merge_cells_in_row(table, row_idx, 0, 3)
                self._set_cell_text(table.cell(row_idx, 0), sec_def.title, bold=True, size=10)
                self._set_cell_shading(table.cell(row_idx, 0), "F1F5F9")
                row_idx += 1

                # 表头 (4列: 巡检内容, 巡检结果, 巡检内容, 巡检结果)
                headers = ["巡检内容", "巡检结果", "巡检内容", "巡检结果"]
                for c, h in enumerate(headers):
                    self._set_cell_text(table.cell(row_idx, c), h, bold=True, size=9,
                                        alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    self._set_cell_shading(table.cell(row_idx, c), "F8FAFC")
                row_idx += 1

                # 2列布局
                items = sec_def.check_items
                for i in range(0, len(items), 2):
                    # 左列
                    self._set_cell_text(table.cell(row_idx, 0), items[i].name, size=9)
                    actual_l = result_map.get(items[i].name)
                    if actual_l:
                        self._set_cell_text(table.cell(row_idx, 1),
                                            f"■ {self._get_status_label(actual_l.status)} □ 异常", size=9)
                    else:
                        self._set_cell_text(table.cell(row_idx, 1), "■ 正常 □ 异常", size=9)

                    # 右列
                    if i + 1 < len(items):
                        self._set_cell_text(table.cell(row_idx, 2), items[i + 1].name, size=9)
                        actual_r = result_map.get(items[i + 1].name)
                        if actual_r:
                            self._set_cell_text(table.cell(row_idx, 3),
                                                f"■ {self._get_status_label(actual_r.status)} □ 异常", size=9)
                        else:
                            self._set_cell_text(table.cell(row_idx, 3), "■ 正常 □ 异常", size=9)
                    row_idx += 1

            elif sec_def.section_type == "blade_slots":
                # 刀片服务器槽位
                self._merge_cells_in_row(table, row_idx, 0, 3)
                self._set_cell_text(table.cell(row_idx, 0), sec_def.title, bold=True, size=10)
                self._set_cell_shading(table.cell(row_idx, 0), "F1F5F9")
                row_idx += 1

                # 表头
                slot_headers = ["", "槽位", "NAME", "用途（主机或IP）", "状态"]
                # 需要5列，但表格是4列，合并方式调整
                self._set_cell_text(table.cell(row_idx, 0), "槽位", bold=True, size=9,
                                    alignment=WD_ALIGN_PARAGRAPH.CENTER)
                self._set_cell_text(table.cell(row_idx, 1), "NAME", bold=True, size=9,
                                    alignment=WD_ALIGN_PARAGRAPH.CENTER)
                self._set_cell_text(table.cell(row_idx, 2), "用途（主机或IP）", bold=True, size=9,
                                    alignment=WD_ALIGN_PARAGRAPH.CENTER)
                self._set_cell_text(table.cell(row_idx, 3), "状态", bold=True, size=9,
                                    alignment=WD_ALIGN_PARAGRAPH.CENTER)
                self._set_cell_shading(table.cell(row_idx, 0), "F8FAFC")
                self._set_cell_shading(table.cell(row_idx, 1), "F8FAFC")
                self._set_cell_shading(table.cell(row_idx, 2), "F8FAFC")
                self._set_cell_shading(table.cell(row_idx, 3), "F8FAFC")
                row_idx += 1

                # 刀片槽位 (从 extra_data 读取或使用默认)
                blade_slots = device.extra_info.get("blade_slots", [
                    {"slot": "1", "name": "", "purpose": "", "status": "正常 异常"},
                    {"slot": "2", "name": "", "purpose": "", "status": "正常 异常"},
                    {"slot": "3", "name": "", "purpose": "", "status": "正常 异常"},
                    {"slot": "11", "name": "", "purpose": "", "status": "正常 异常"},
                ])
                for slot in blade_slots[:4]:
                    self._set_cell_text(table.cell(row_idx, 0), slot.get("slot", ""), size=9)
                    self._set_cell_text(table.cell(row_idx, 1), slot.get("name", ""), size=9)
                    self._set_cell_text(table.cell(row_idx, 2), slot.get("purpose", ""), size=9)
                    self._set_cell_text(table.cell(row_idx, 3), slot.get("status", "正常 异常"), size=9)
                    row_idx += 1

    def _generate_device_table_generic(self, device_result: DeviceInspectionResult):
        """通用设备巡检表 (无模板匹配时)"""
        device = device_result.device

        # 基本信息
        self._add_paragraph(f"设备类型: {device.device_type.value}", size=10)
        self._add_paragraph(f"品牌型号: {device.brand} {device.model}", size=10)
        if device.ip:
            self._add_paragraph(f"管理IP: {device.ip}", size=10)
        if device.serial_number:
            self._add_paragraph(f"序列号: {device.serial_number}", size=10)

        self.doc.add_paragraph()

        # 巡检结果表
        all_items = []
        for sec in device_result.sections:
            all_items.extend(sec.items)

        if not all_items:
            self._add_paragraph("暂无巡检数据", color=self.COLOR_MUTED)
            return

        headers = ["序号", "检查项", "检查结果", "状态", "建议"]
        table = self.doc.add_table(rows=1 + len(all_items), cols=len(headers))
        table.style = 'Table Grid'

        for c, h in enumerate(headers):
            self._set_cell_text(table.cell(0, c), h, bold=True, size=9,
                                alignment=WD_ALIGN_PARAGRAPH.CENTER)
            self._set_cell_shading(table.cell(0, c), "F1F5F9")

        for r, item in enumerate(all_items):
            self._set_cell_text(table.cell(r + 1, 0), str(r + 1), size=8)
            self._set_cell_text(table.cell(r + 1, 1), item.name, size=8)
            self._set_cell_text(table.cell(r + 1, 2), item.raw_value or item.result, size=8)
            self._set_cell_text(table.cell(r + 1, 3), self._get_status_label(item.status), size=8)
            self._set_cell_text(table.cell(r + 1, 4), item.suggestion or "-", size=8)

    def _get_info_values(self, device: DeviceInfo, sec_def: SectionDef) -> Dict[str, str]:
        """从设备信息中提取基本信息值"""
        values = {}
        hw = device.hardware_config or {}

        # 基本映射
        field_map = {
            "主机型号": device.model,
            "设备型号": device.model,
            "序列号": device.serial_number,
            "系统版本": device.os_version,
            "固件版本": device.os_version,
            "管理IP": device.ip,
            "操作系统版本": device.os_version,
            "vCenter版本": device.os_version,
            "ESXi版本": device.os_version,
            "备份软件": device.extra_info.get("backup_software", ""),
            "服务器型号": device.model,
            "操作系统": device.os_version,
        }

        # 硬件配置映射
        hw_map = {
            "接口数量": hw.get("interfaces", ""),
            "电源数量": hw.get("power_supplies", ""),
            "设备运行环境": hw.get("environment", "√良好 □一般 □恶劣"),
            "模块信息": hw.get("modules", ""),
            "模块": hw.get("modules", ""),
            "端口数量": hw.get("ports", ""),
            "CPU/内存配置": hw.get("cpu_memory", ""),
            "分区信息": hw.get("partition", ""),
            "集群名称": hw.get("cluster_name", ""),
            "主机数量": hw.get("host_count", ""),
            "虚拟机数量": hw.get("vm_count", ""),
            "管理IP": device.ip,
        }

        for field_label, _ in sec_def.info_fields:
            if not field_label:
                continue
            if field_label in field_map:
                values[field_label] = field_map[field_label] or ""
            elif field_label in hw_map:
                values[field_label] = hw_map[field_label] or ""

        # 子字段
        for _, placeholder in sec_def.info_fields:
            if placeholder and placeholder in hw_map:
                values[placeholder] = hw_map[placeholder] or ""

        # 额外信息
        for k, v in device.extra_info.items():
            if k not in values:
                values[k] = str(v)

        return values

    def _get_status_label(self, status: str) -> str:
        """获取状态标签"""
        labels = {
            "ok": "正常",
            "warning": "警告",
            "critical": "严重",
            "error": "错误",
            "skipped": "跳过",
        }
        return labels.get(status, status or "正常")

    # ============================================================
    # 数据库巡检附录
    # ============================================================

    def _generate_database_sections(self):
        """生成数据库巡检报告段落"""
        db_results = [dr for dr in self.data.device_results
                      if dr.device.device_type.value == "database"]
        if not db_results:
            return

        self._add_heading("数据库系统巡检报告", level=1)
        self._add_paragraph(
            "本次数据库常规检查的数据收集主要集中在巡检周期内，我们尽可能把重要的信息收集起来进行分析，"
            "此次常规检查主要是根据维保合同设备清单，按照巡检计划到达客户现场，"
            "进行数据库和备份软件的巡检和优化调整。",
            size=10
        )

        for dr in db_results:
            self._add_heading(f"{dr.device.name}-{dr.device.ip}", level=2)
            # 数据库使用通用表格
            self._generate_device_table_generic(dr)

        self.doc.add_page_break()

    # ============================================================
    # 问题汇总
    # ============================================================

    def _generate_issues_summary(self):
        """生成问题汇总"""
        if not self.data.issues:
            return

        self._add_heading("主要问题汇总", level=1)
        self._add_paragraph(
            f"共发现 {len(self.data.issues)} 个问题（严重: {self.data.critical_count}, 警告: {self.data.warning_count}）",
            size=10, color=self.COLOR_MUTED
        )

        critical_issues = [i for i in self.data.issues if i.get("status") == "critical"]
        warning_issues = [i for i in self.data.issues if i.get("status") == "warning"]

        if critical_issues:
            self._add_heading("严重问题", level=2)
            headers = ["对象", "检查项", "当前值", "阈值", "建议"]
            table = self.doc.add_table(rows=1 + len(critical_issues), cols=len(headers))
            table.style = 'Table Grid'
            for c, h in enumerate(headers):
                self._set_cell_text(table.cell(0, c), h, bold=True, size=9,
                                    alignment=WD_ALIGN_PARAGRAPH.CENTER)
                self._set_cell_shading(table.cell(0, c), "F1F5F9")
            for r, issue in enumerate(critical_issues):
                self._set_cell_text(table.cell(r + 1, 0), issue.get("target_name", ""), size=8)
                self._set_cell_text(table.cell(r + 1, 1), issue.get("item_name", ""), size=8)
                self._set_cell_text(table.cell(r + 1, 2), issue.get("value", ""), size=8)
                self._set_cell_text(table.cell(r + 1, 3), issue.get("threshold", ""), size=8)
                self._set_cell_text(table.cell(r + 1, 4), issue.get("suggestion", ""), size=8)
            self.doc.add_paragraph()

        if warning_issues:
            self._add_heading("警告问题", level=2)
            headers = ["对象", "检查项", "当前值", "阈值", "建议"]
            table = self.doc.add_table(rows=1 + len(warning_issues), cols=len(headers))
            table.style = 'Table Grid'
            for c, h in enumerate(headers):
                self._set_cell_text(table.cell(0, c), h, bold=True, size=9,
                                    alignment=WD_ALIGN_PARAGRAPH.CENTER)
                self._set_cell_shading(table.cell(0, c), "F1F5F9")
            for r, issue in enumerate(warning_issues):
                self._set_cell_text(table.cell(r + 1, 0), issue.get("target_name", ""), size=8)
                self._set_cell_text(table.cell(r + 1, 1), issue.get("item_name", ""), size=8)
                self._set_cell_text(table.cell(r + 1, 2), issue.get("value", ""), size=8)
                self._set_cell_text(table.cell(r + 1, 3), issue.get("threshold", ""), size=8)
                self._set_cell_text(table.cell(r + 1, 4), issue.get("suggestion", ""), size=8)

        self.doc.add_page_break()

    # ============================================================
    # 附录
    # ============================================================

    def _generate_appendix(self):
        """生成附录"""
        self._add_heading("附录", level=1)
        self._add_heading("报告信息", level=2)

        app_headers = ["项目", "内容"]
        app_rows = [
            ["报告编号", self.data.report_id],
            ["任务名称", self.data.task_name],
            ["任务ID", self.data.task_id],
            ["生成时间", self.data.generated_at],
            ["报告格式", "DOCX (可编辑)"],
            ["巡检对象数", str(self.data.target_count)],
            ["总检查项", str(self.data.total_items)],
            ["健康度评分", f"{self.data.health_score}/100"],
        ]
        table = self.doc.add_table(rows=len(app_rows), cols=2)
        table.style = 'Table Grid'
        for r, (k, v) in enumerate(app_rows):
            self._set_cell_text(table.cell(r, 0), k, bold=True, size=9)
            self._set_cell_text(table.cell(r, 1), v, size=9)

        self.doc.add_paragraph()
        self._add_paragraph(
            f"— 本报告由 {self.data.config.platform_name} 自动生成 —",
            size=8, color=self.COLOR_MUTED,
            alignment=WD_ALIGN_PARAGRAPH.CENTER
        )

    # ============================================================
    # 主生成流程
    # ============================================================

    def generate(self, output_path: str = None) -> str:
        """生成DOCX报告"""
        if not output_path:
            output_path = os.path.join(OUTPUT_DIR, f"{self.data.report_id}.docx")

        # 1. 封面
        if self.data.config.include_cover:
            self._generate_cover()

        # 2. 文档控制
        self._generate_doc_control()

        # 3. 目录
        if self.data.config.include_toc:
            self._generate_toc()

        # 4. 巡检任务信息
        self._generate_task_info()

        # 5. 设备巡检详细结果 (按报告分区)
        if self.data.device_results:
            # 按设备类型分组
            network_devices = [dr for dr in self.data.device_results
                               if dr.device.device_type.value in ("router", "switch", "firewall",
                                                                   "load_balancer", "ips", "waf",
                                                                   "vpn", "bastion", "netgap")]
            server_devices = [dr for dr in self.data.device_results
                              if dr.device.device_type.value in ("server", "ibm_minicomputer",
                                                                  "blade_center")]
            storage_devices = [dr for dr in self.data.device_results
                               if dr.device.device_type.value in ("storage", "san_switch", "tape_library")]
            vm_devices = [dr for dr in self.data.device_results
                          if dr.device.device_type.value in ("vmware",)]
            backup_devices = [dr for dr in self.data.device_results
                              if dr.device.device_type.value in ("backup",)]
            db_devices = [dr for dr in self.data.device_results
                          if dr.device.device_type.value in ("database",)]

            if network_devices:
                self._add_heading("市民卡业务巡检报告", level=1)
                for dr in network_devices:
                    self._generate_device_section(dr)
                self.doc.add_page_break()

            if server_devices:
                self._add_heading("主机巡检报告", level=1)
                for dr in server_devices:
                    self._generate_device_section(dr)
                self.doc.add_page_break()

            if storage_devices:
                self._add_heading("存储巡检报告", level=1)
                for dr in storage_devices:
                    self._generate_device_section(dr)
                self.doc.add_page_break()

            if vm_devices:
                self._add_heading("虚拟化巡检报告", level=1)
                for dr in vm_devices:
                    self._generate_device_section(dr)
                self.doc.add_page_break()

            if backup_devices:
                self._add_heading("备份系统巡检报告", level=1)
                for dr in backup_devices:
                    self._generate_device_section(dr)
                self.doc.add_page_break()

            if db_devices:
                self._generate_database_sections()
        else:
            # 无设备结果时，使用旧的通用格式
            self._generate_issues_summary()

        # 6. 附录
        if self.data.config.include_appendix:
            self._generate_appendix()

        # 保存
        self.doc.save(output_path)
        return output_path


def generate_sample_report() -> str:
    """生成示例报告用于测试"""
    from . import (InspectionReportData, ReportConfig, DeviceInspectionResult,
                   DeviceInfo, InspectionSection, InspectionCheckItem, DeviceType)

    # 构建设备巡检结果
    device_results = []

    # 路由器
    router = DeviceInfo(
        name="生产边界路由器",
        ip="10.88.1.1",
        device_type=DeviceType.ROUTER,
        brand="H3C",
        model="SR6608-X",
        serial_number="210231A1U5B156000186",
        os_version="Comware Software, Version 5.20.106, Release 3303P16",
        hardware_config={
            "interfaces": "60千兆电口，12千兆光口",
            "power_supplies": "2",
            "environment": "√良好 □一般 □恶劣",
        }
    )
    router_result = DeviceInspectionResult(
        device=router,
        sections=[
            InspectionSection(title="设备运行状态检查", items=[
                InspectionCheckItem(name="设备指示灯", result="正常", status="ok", raw_value="无红灯"),
                InspectionCheckItem(name="当前及历史CPU使用率", result="正常", status="ok", raw_value="峰值15%"),
                InspectionCheckItem(name="内存使用率", result="正常", status="ok", raw_value="62%"),
                InspectionCheckItem(name="风扇状态", result="正常", status="ok"),
                InspectionCheckItem(name="电源状态", result="正常", status="ok"),
                InspectionCheckItem(name="模块温度", result="正常", status="ok", raw_value="35℃"),
                InspectionCheckItem(name="端口状态", result="正常", status="ok"),
                InspectionCheckItem(name="日志检查", result="正常", status="ok"),
            ]),
        ],
        overall_status="ok",
    )
    device_results.append(router_result)

    # 交换机
    switch = DeviceInfo(
        name="内网核心交换机",
        ip="10.88.41.1",
        device_type=DeviceType.SWITCH,
        brand="H3C",
        model="S10508-V",
        serial_number="210235A0JAX157000012",
        os_version="H3C Comware Software, Version 7.1.045, Release 7150P02",
        hardware_config={
            "interfaces": "56千兆电口，24千兆光口，12万兆光口",
            "power_supplies": "2",
        }
    )
    switch_result = DeviceInspectionResult(
        device=switch,
        sections=[
            InspectionSection(title="设备运行状态检查", items=[
                InspectionCheckItem(name="设备指示灯", result="正常", status="ok"),
                InspectionCheckItem(name="当前及历史CPU使用率", result="正常", status="ok", raw_value="22%"),
                InspectionCheckItem(name="内存使用率", result="正常", status="ok", raw_value="45%"),
            ]),
        ],
        overall_status="ok",
    )
    device_results.append(switch_result)

    # 防火墙
    fw = DeviceInfo(
        name="生产边界防火墙主",
        ip="10.88.40.1",
        device_type=DeviceType.FIREWALL,
        brand="",
        model="",
        hardware_config={
            "interfaces": "4个千兆电口,4个千兆光口，2个万兆光口",
            "power_supplies": "2",
        }
    )
    fw_result = DeviceInspectionResult(
        device=fw,
        sections=[
            InspectionSection(title="设备运行状态检查", items=[
                InspectionCheckItem(name="设备面板指示灯状态检查", result="正常", status="ok"),
                InspectionCheckItem(name="设备性能状态检查", result="正常", status="ok"),
                InspectionCheckItem(name="设备双机热备状态检查", result="正常", status="ok"),
            ]),
        ],
        overall_status="ok",
    )
    device_results.append(fw_result)

    config = ReportConfig(
        title="人社市民卡及五险硬件系统维保采购项目",
        subtitle="巡检服务报告",
        company_name="义乌市人力资源和社会保障信息中心",
        client_contact="丁晓波",
        client_phone="13819959652",
        contract_number="SJZJZC2023230GK",
        inspection_period="月度",
        inspection_date_range="2024.09.23-9.29",
        service_team="天健技术服务团队",
        engineers=[
            {"role": "主机和数据库MA", "name": "郎军伟/章伟东", "phone": "400-9900-806"},
            {"role": "网络MA", "name": "范俊涛", "phone": "400-9900-806"},
        ],
        report_authors=["范俊涛", "郎军伟", "章伟东"],
        report_reviewer="杜婉鹏",
        author_date="2024.09.23",
        reviewer_date="2024.09.29",
        service_content="根据维保合同设备清单，按照月度巡检计划到达客户现场，进行网络、主机、存储和数据库的巡检和优化调整",
    )

    data = InspectionReportData(
        task_name="月度巡检",
        task_id="task-monthly-001",
        target_count=3,
        health_score=95,
        total_items=14,
        ok_count=14,
        warning_count=0,
        critical_count=0,
        summary="本次巡检所有设备运行状态正常，未发现严重问题。",
        device_results=device_results,
        config=config,
    )

    gen = DocxGenerator(data)
    return gen.generate()
