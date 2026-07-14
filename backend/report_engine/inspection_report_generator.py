"""
Acdante ITOps - 巡检报告Word模板生成器
基于义乌社保巡检报告风格，生成标准化巡检报告
"""

import os
import time
import logging
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)

OUTPUT_DIR = os.environ.get("ITOPS_REPORT_DIR", os.path.join(os.path.dirname(__file__), "..", "data", "reports"))


def generate_inspection_report(
    task_name: str,
    task_id: str,
    customer_name: str,
    inspection_date: str,
    inspection_cycle: str,
    engineers: List[Dict],
    devices: List[Dict],
    results: List[Dict],
    issues: List[Dict],
    recommendations: List[str],
    output_format: str = "docx",
) -> Dict:
    """
    生成标准化巡检报告
    
    参数:
        task_name: 巡检任务名称
        task_id: 任务ID
        customer_name: 客户名称
        inspection_date: 巡检日期
        inspection_cycle: 巡检周期
        engineers: 巡检工程师列表 [{"name": "", "role": "", "phone": ""}]
        devices: 设备清单 [{"seq": 1, "type": "网络", "device_type": "路由器", "brand": "H3C", "model": "SR6608-X", "quantity": 1, "serial": "", "ip": "", "remark": ""}]
        results: 巡检结果 [{"device_name": "", "device_ip": "", "check_items": [{"name": "", "command": "", "status": "normal/abnormal", "result": ""}]}]
        issues: 问题列表 [{"device_name": "", "ip": "", "description": "", "severity": "高/中/低", "suggestion": ""}]
        recommendations: 维护建议列表
        output_format: 输出格式 (docx/html)
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    report_id = f"rpt-{int(time.time() * 1000) % 1000000:06d}"

    if output_format == "docx":
        output_path = os.path.join(OUTPUT_DIR, f"{report_id}.docx")
        _generate_docx_report(
            output_path, task_name, customer_name, inspection_date,
            inspection_cycle, engineers, devices, results, issues, recommendations
        )
    else:
        output_path = os.path.join(OUTPUT_DIR, f"{report_id}.html")
        _generate_html_report(
            output_path, task_name, customer_name, inspection_date,
            inspection_cycle, engineers, devices, results, issues, recommendations
        )

    # 统计
    total_devices = len(devices)
    total_items = sum(len(r.get("check_items", [])) for r in results)
    ok_count = sum(1 for r in results for item in r.get("check_items", []) if item.get("status") == "normal")
    abnormal_count = sum(1 for r in results for item in r.get("check_items", []) if item.get("status") == "abnormal")
    health_score = int(ok_count / total_items * 100) if total_items > 0 else 100

    return {
        "id": report_id,
        "task_id": task_id,
        "task_name": task_name,
        "format": output_format,
        "health_score": health_score,
        "total_devices": total_devices,
        "total_items": total_items,
        "ok_count": ok_count,
        "abnormal_count": abnormal_count,
        "issue_count": len(issues),
        "generated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "file_path": output_path,
        "download_url": f"/api/v1/reports/{report_id}/download",
    }


def _generate_docx_report(
    output_path, task_name, customer_name, inspection_date,
    inspection_cycle, engineers, devices, results, issues, recommendations
):
    """生成DOCX格式巡检报告"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    doc = Document()
    FONT_NAME = '宋体'
    FONT_COLOR = RGBColor(0, 0, 0)

    def set_cell(cell, text, bold=False, size=None, align='left', bg=None):
        cell.text = ''
        p = cell.paragraphs[0]
        if align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(text))
        run.font.name = FONT_NAME
        run.font.color.rgb = FONT_COLOR
        run.font.size = size or Pt(10.5)
        run.font.bold = bold
        run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        if bg:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}"/>')
            cell._tc.get_or_add_tcPr().append(shading)

    def merge_cells(table, r1, c1, r2, c2):
        table.cell(r1, c1).merge(table.cell(r2, c2))

    # 封面
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(customer_name)
    run.font.name = '黑体'
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{inspection_date[:4]}年M{inspection_date[5:7]}月度巡检服务报告')
    run.font.name = '黑体'
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_page_break()

    # 巡检任务信息
    h = doc.add_heading('巡检服务报告', level=1)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    table = doc.add_table(rows=5 + len(engineers), cols=6)
    table.style = 'Table Grid'

    merge_cells(table, 0, 0, 0, 5)
    set_cell(table.rows[0].cells[0], '巡检任务信息', bold=True, size=Pt(12), align='center', bg='2E86C1')
    table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    set_cell(table.rows[1].cells[0], '用户单位', bold=True, bg='D6EAF8')
    merge_cells(table, 1, 1, 1, 5)
    set_cell(table.rows[1].cells[1], customer_name)

    set_cell(table.rows[2].cells[0], '巡检时间', bold=True, bg='D6EAF8')
    set_cell(table.rows[2].cells[1], inspection_date)
    set_cell(table.rows[2].cells[2], '')
    set_cell(table.rows[2].cells[3], '巡检周期', bold=True, bg='D6EAF8')
    merge_cells(table, 2, 4, 2, 5)
    set_cell(table.rows[2].cells[4], inspection_cycle)

    set_cell(table.rows[3].cells[0], '巡检工程师', bold=True, bg='D6EAF8')
    set_cell(table.rows[3].cells[1], '职责', bold=True, bg='D6EAF8', align='center')
    merge_cells(table, 3, 2, 3, 3)
    set_cell(table.rows[3].cells[2], '姓名', bold=True, bg='D6EAF8', align='center')
    merge_cells(table, 3, 4, 3, 5)
    set_cell(table.rows[3].cells[4], '联系方式', bold=True, bg='D6EAF8', align='center')

    for i, eng in enumerate(engineers):
        row = table.rows[4 + i]
        set_cell(row.cells[0], '')
        set_cell(row.cells[1], eng.get('role', ''))
        merge_cells(row, 2, row, 3)
        set_cell(row.cells[2], eng.get('name', ''))
        merge_cells(row, 4, row, 5)
        set_cell(row.cells[4], eng.get('phone', ''))

    # 设备清单
    h = doc.add_heading('一、硬件系统巡检清单', level=2)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    table = doc.add_table(rows=1 + len(devices), cols=8)
    table.style = 'Table Grid'
    headers = ['序号', '类型', '设备类型', '品牌型号', '数量', '序列号', 'IP地址', '备注']
    for i, h_text in enumerate(headers):
        set_cell(table.rows[0].cells[i], h_text, bold=True, bg='2E86C1', align='center', size=Pt(9))
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for r_idx, dev in enumerate(devices):
        row = table.rows[1 + r_idx]
        set_cell(row.cells[0], str(dev.get('seq', r_idx + 1)), size=Pt(9), align='center')
        set_cell(row.cells[1], dev.get('type', ''), size=Pt(9), align='center')
        set_cell(row.cells[2], dev.get('device_type', ''), size=Pt(9), align='center')
        set_cell(row.cells[3], f"{dev.get('brand', '')} {dev.get('model', '')}", size=Pt(9), align='center')
        set_cell(row.cells[4], str(dev.get('quantity', 1)), size=Pt(9), align='center')
        set_cell(row.cells[5], dev.get('serial', ''), size=Pt(9), align='center')
        set_cell(row.cells[6], dev.get('ip', ''), size=Pt(9), align='center')
        set_cell(row.cells[7], dev.get('remark', ''), size=Pt(9), align='center')

    doc.add_paragraph()

    # 巡检总结
    h = doc.add_heading('二、巡检总结', level=2)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # 统计
    total_items = sum(len(r.get("check_items", [])) for r in results)
    ok_count = sum(1 for r in results for item in r.get("check_items", []) if item.get("status") == "normal")
    abnormal_count = total_items - ok_count

    p = doc.add_paragraph()
    run = p.add_run(f'本次巡检共检查设备 {len(devices)} 台，巡检项目 {total_items} 项。其中正常 {ok_count} 项，异常 {abnormal_count} 项。')
    run.font.name = FONT_NAME
    run.font.size = Pt(10.5)

    # 问题汇总
    if issues:
        h = doc.add_heading('主要问题汇总', level=3)
        for run in h.runs:
            run.font.name = '黑体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        table = doc.add_table(rows=1 + len(issues), cols=6)
        table.style = 'Table Grid'
        issue_headers = ['序号', '设备名称', 'IP地址', '问题描述', '严重程度', '处理建议']
        for i, h_text in enumerate(issue_headers):
            set_cell(table.rows[0].cells[i], h_text, bold=True, bg='E74C3C', align='center', size=Pt(9))
            table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        for r_idx, issue in enumerate(issues):
            row = table.rows[1 + r_idx]
            set_cell(row.cells[0], str(r_idx + 1), size=Pt(9), align='center')
            set_cell(row.cells[1], issue.get('device_name', ''), size=Pt(9))
            set_cell(row.cells[2], issue.get('ip', ''), size=Pt(9))
            set_cell(row.cells[3], issue.get('description', ''), size=Pt(9))
            severity = issue.get('severity', '中')
            bg = 'FADBD8' if severity == '高' else ('FEF9E7' if severity == '中' else 'D5F5E3')
            set_cell(row.cells[4], severity, size=Pt(9), align='center', bg=bg)
            set_cell(row.cells[5], issue.get('suggestion', ''), size=Pt(9))

    # 维护建议
    if recommendations:
        doc.add_paragraph()
        h = doc.add_heading('维护建议', level=3)
        for run in h.runs:
            run.font.name = '黑体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        for i, rec in enumerate(recommendations):
            p = doc.add_paragraph()
            run = p.add_run(f'（{i+1}）{rec}')
            run.font.name = FONT_NAME
            run.font.size = Pt(10.5)

    doc.add_page_break()

    # 设备详细巡检
    h = doc.add_heading('三、设备详细巡检信息', level=2)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    for device_result in results:
        device_name = device_result.get('device_name', '未知设备')
        device_ip = device_result.get('device_ip', '')
        check_items = device_result.get('check_items', [])

        h = doc.add_heading(f'{device_name}-{device_ip}', level=3)
        for run in h.runs:
            run.font.name = '黑体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        table = doc.add_table(rows=2 + len(check_items), cols=4)
        table.style = 'Table Grid'

        merge_cells(table, 0, 0, 0, 3)
        set_cell(table.rows[0].cells[0], '设备运行状态检查', bold=True, bg='2E86C1', align='center')
        table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        set_cell(table.rows[1].cells[0], '检查内容', bold=True, bg='D6EAF8', align='center')
        set_cell(table.rows[1].cells[1], '检查操作', bold=True, bg='D6EAF8', align='center')
        set_cell(table.rows[1].cells[2], '巡检结果', bold=True, bg='D6EAF8', align='center')
        set_cell(table.rows[1].cells[3], '结果说明', bold=True, bg='D6EAF8', align='center')

        for i, item in enumerate(check_items):
            row = table.rows[2 + i]
            set_cell(row.cells[0], item.get('name', ''), bold=True, size=Pt(9))
            set_cell(row.cells[1], item.get('command', ''), size=Pt(9))

            status = item.get('status', 'normal')
            if status == 'normal':
                set_cell(row.cells[2], '√ 正常', bold=True, size=Pt(9), align='center', bg='D5F5E3')
            else:
                set_cell(row.cells[2], '× 异常', bold=True, size=Pt(9), align='center', bg='FADBD8')

            set_cell(row.cells[3], item.get('result', ''), size=Pt(9))

        doc.add_paragraph()

    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    doc.save(output_path)
    logger.info(f"巡检报告已生成: {output_path}")


def _generate_html_report(
    output_path, task_name, customer_name, inspection_date,
    inspection_cycle, engineers, devices, results, issues, recommendations
):
    """生成HTML格式巡检报告"""
    # 统计
    total_items = sum(len(r.get("check_items", [])) for r in results)
    ok_count = sum(1 for r in results for item in r.get("check_items", []) if item.get("status") == "normal")
    abnormal_count = total_items - ok_count
    health_score = int(ok_count / total_items * 100) if total_items > 0 else 100

    # 设备清单行
    device_rows = ""
    for dev in devices:
        device_rows += f"""<tr>
            <td>{dev.get('seq', '')}</td><td>{dev.get('type', '')}</td>
            <td>{dev.get('device_type', '')}</td><td>{dev.get('brand', '')} {dev.get('model', '')}</td>
            <td>{dev.get('quantity', 1)}</td><td>{dev.get('serial', '')}</td>
            <td>{dev.get('ip', '')}</td><td>{dev.get('remark', '')}</td>
        </tr>"""

    # 问题行
    issue_rows = ""
    for i, issue in enumerate(issues):
        severity = issue.get('severity', '中')
        color = '#e74c3c' if severity == '高' else ('#f39c12' if severity == '中' else '#27ae60')
        issue_rows += f"""<tr>
            <td>{i+1}</td><td>{issue.get('device_name', '')}</td>
            <td>{issue.get('ip', '')}</td><td>{issue.get('description', '')}</td>
            <td style="color:{color};font-weight:bold">{severity}</td>
            <td>{issue.get('suggestion', '')}</td>
        </tr>"""

    # 设备详情
    device_details = ""
    for device_result in results:
        device_name = device_result.get('device_name', '未知设备')
        device_ip = device_result.get('device_ip', '')
        check_items = device_result.get('check_items', [])

        items_html = ""
        for item in check_items:
            status = item.get('status', 'normal')
            if status == 'normal':
                status_html = '<span style="color:#27ae60;font-weight:bold">√ 正常</span>'
                bg = '#d5f5e3'
            else:
                status_html = '<span style="color:#e74c3c;font-weight:bold">× 异常</span>'
                bg = '#fadbd8'

            items_html += f"""<tr>
                <td style="font-weight:bold">{item.get('name', '')}</td>
                <td>{item.get('command', '')}</td>
                <td style="background:{bg};text-align:center">{status_html}</td>
                <td>{item.get('result', '')}</td>
            </tr>"""

        device_details += f"""
        <h3>{device_name}-{device_ip}</h3>
        <table class="detail-table">
            <thead><tr><th colspan="4" style="background:#2E86C1;color:white;text-align:center">设备运行状态检查</th></tr>
            <tr><th>检查内容</th><th>检查操作</th><th>巡检结果</th><th>结果说明</th></tr></thead>
            <tbody>{items_html}</tbody>
        </table>"""

    # 建议
    recommendations_html = ""
    for i, rec in enumerate(recommendations):
        recommendations_html += f"<p>（{i+1}）{rec}</p>"

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>{task_name} - 巡检报告</title>
<style>
body {{ font-family: 宋体, SimSun, serif; margin: 0; padding: 20px; color: #333; font-size: 14px; }}
.container {{ max-width: 1200px; margin: 0 auto; }}
h1 {{ color: #003366; font-family: 黑体, SimHei, serif; border-bottom: 3px solid #2E86C1; padding-bottom: 10px; }}
h2 {{ color: #003366; font-family: 黑体, SimHei, serif; border-bottom: 2px solid #2E86C1; padding-bottom: 8px; }}
h3 {{ color: #2E86C1; font-family: 黑体, SimHei, serif; }}
table {{ width: 100%; border-collapse: collapse; margin: 16px 0; font-size: 13px; }}
th, td {{ border: 1px solid #bdc3c7; padding: 8px 12px; text-align: left; }}
th {{ background: #2E86C1; color: white; text-align: center; }}
.info-table th {{ background: #D6EAF8; color: #333; width: 120px; }}
.detail-table th {{ background: #D6EAF8; color: #333; }}
.stats {{ display: flex; gap: 20px; margin: 20px 0; }}
.stat {{ flex: 1; padding: 20px; border-radius: 8px; text-align: center; }}
.stat-ok {{ background: #d5f5e3; color: #27ae60; }}
.stat-abnormal {{ background: #fadbd8; color: #e74c3c; }}
.stat-score {{ background: #d6eaf8; color: #2e86c1; }}
.stat h3 {{ margin: 0; font-size: 36px; }}
.stat p {{ margin: 5px 0 0; font-size: 14px; }}
.footer {{ margin-top: 40px; padding-top: 20px; border-top: 2px solid #2E86C1; text-align: center; color: #7f8c8d; font-size: 12px; }}
</style>
</head>
<body>
<div class="container">
    <h1>{task_name}</h1>
    <p><strong>客户:</strong> {customer_name} | <strong>日期:</strong> {inspection_date} | <strong>周期:</strong> {inspection_cycle}</p>

    <div class="stats">
        <div class="stat stat-ok"><h3>{ok_count}</h3><p>正常项</p></div>
        <div class="stat stat-abnormal"><h3>{abnormal_count}</h3><p>异常项</p></div>
        <div class="stat stat-score"><h3>{health_score}%</h3><p>健康度</p></div>
    </div>

    <h2>一、设备清单</h2>
    <table><thead><tr><th>序号</th><th>类型</th><th>设备类型</th><th>品牌型号</th><th>数量</th><th>序列号</th><th>IP地址</th><th>备注</th></tr></thead>
    <tbody>{device_rows}</tbody></table>

    <h2>二、巡检总结</h2>
    <p>本次巡检共检查设备 {len(devices)} 台，巡检项目 {total_items} 项。正常 {ok_count} 项，异常 {abnormal_count} 项。</p>

    {"<h3>主要问题汇总</h3><table><thead><tr><th>序号</th><th>设备名称</th><th>IP</th><th>问题描述</th><th>严重程度</th><th>处理建议</th></tr></thead><tbody>" + issue_rows + "</tbody></table>" if issues else ""}

    {"<h3>维护建议</h3>" + recommendations_html if recommendations else ""}

    <h2>三、设备详细巡检信息</h2>
    {device_details}

    <div class="footer">
        <p>Acdante ITOps Inspection Platform | {inspection_date}</p>
    </div>
</div>
</body>
</html>"""

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html)

    logger.info(f"巡检报告已生成: {output_path}")
