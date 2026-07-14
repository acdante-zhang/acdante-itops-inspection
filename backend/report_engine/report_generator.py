"""
Acdante ITOps - 巡检报告生成引擎
支持 HTML / DOCX / PDF 格式
"""

import os
import time
import logging
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)

OUTPUT_DIR = os.environ.get("ITOPS_REPORT_DIR", os.path.join(os.path.dirname(__file__), "..", "data", "reports"))


def ensure_output_dir():
    os.makedirs(OUTPUT_DIR, exist_ok=True)


def generate_report(
    task_name: str,
    task_id: str,
    targets: List[Dict],
    results: List[Dict],
    format: str = "html",
    config: Dict = None,
) -> Dict:
    """生成巡检报告"""
    ensure_output_dir()

    report_id = f"rpt-{int(time.time() * 1000) % 1000000:06d}"
    timestamp = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())

    # 统计
    total_items = len(results)
    ok_count = sum(1 for r in results if r.get("status") == "ok")
    warning_count = sum(1 for r in results if r.get("status") == "warning")
    critical_count = sum(1 for r in results if r.get("status") == "critical")
    error_count = sum(1 for r in results if r.get("status") == "error")

    # 健康分数
    if total_items > 0:
        health_score = int((ok_count * 100 + warning_count * 60 + error_count * 50) / total_items)
    else:
        health_score = 100

    # 生成摘要
    issues = [r for r in results if r.get("status") in ("warning", "critical")]
    if critical_count > 0:
        summary = f"发现 {critical_count} 个严重问题，{warning_count} 个警告，需立即处理"
    elif warning_count > 0:
        summary = f"发现 {warning_count} 个警告，建议关注"
    else:
        summary = "巡检正常，未发现异常"

    generated_files = {}

    if format in ("html", "all"):
        html_path = os.path.join(OUTPUT_DIR, f"{report_id}.html")
        _generate_html(html_path, task_name, timestamp, health_score,
                      total_items, ok_count, warning_count, critical_count,
                      targets, results, summary)
        generated_files["html"] = html_path

    if format in ("docx", "all"):
        try:
            docx_path = os.path.join(OUTPUT_DIR, f"{report_id}.docx")
            _generate_docx(docx_path, task_name, timestamp, health_score,
                          total_items, ok_count, warning_count, critical_count,
                          targets, results, summary)
            generated_files["docx"] = docx_path
        except Exception as e:
            logger.warning(f"DOCX生成失败: {e}")

    return {
        "id": report_id,
        "task_id": task_id,
        "task_name": task_name,
        "format": format,
        "health_score": health_score,
        "total_items": total_items,
        "ok_count": ok_count,
        "warning_count": warning_count,
        "critical_count": critical_count,
        "summary": summary,
        "generated_at": timestamp,
        "files": generated_files,
        "download_url": f"/api/v1/reports/{report_id}/download",
    }


def _generate_html(path, task_name, timestamp, health_score,
                   total, ok, warn, crit, targets, results, summary):
    """生成HTML报告"""
    status_colors = {"ok": "#22c55e", "warning": "#f59e0b", "critical": "#ef4444", "error": "#ef4444"}
    status_labels = {"ok": "正常", "warning": "警告", "critical": "严重", "error": "错误"}

    rows = ""
    for r in results:
        color = status_colors.get(r.get("status", ""), "#94a3b8")
        label = status_labels.get(r.get("status", ""), "未知")
        rows += f"""<tr>
            <td>{r.get('target_name', '')}</td>
            <td>{r.get('item_name', '')}</td>
            <td>{r.get('category', '')}</td>
            <td style="color:{color};font-weight:600">{label}</td>
            <td style="font-family:monospace;font-size:12px">{str(r.get('raw_value', ''))[:100]}</td>
            <td>{r.get('suggestion', '')}</td>
        </tr>"""

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>{task_name} - 巡检报告</title>
<style>
body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; margin: 0; padding: 20px; background: #f8fafc; color: #334155; }}
.container {{ max-width: 1200px; margin: 0 auto; }}
.header {{ background: linear-gradient(135deg, #0f172a, #1e293b); color: #fff; padding: 32px; border-radius: 12px; margin-bottom: 24px; }}
.header h1 {{ margin: 0 0 8px; font-size: 24px; }}
.header p {{ margin: 0; color: #94a3b8; font-size: 14px; }}
.stats {{ display: flex; gap: 16px; margin: 24px 0; }}
.stat {{ flex: 1; padding: 20px; border-radius: 8px; text-align: center; }}
.stat-ok {{ background: #dcfce7; color: #166534; }}
.stat-warn {{ background: #fef3c7; color: #92400e; }}
.stat-crit {{ background: #fecaca; color: #991b1b; }}
.stat-score {{ background: #f0f9ff; color: #075985; }}
.stat h3 {{ margin: 0; font-size: 32px; }}
.stat p {{ margin: 4px 0 0; font-size: 13px; }}
table {{ width: 100%; border-collapse: collapse; margin: 16px 0; background: #fff; border-radius: 8px; overflow: hidden; box-shadow: 0 1px 3px rgba(0,0,0,.1); }}
th {{ background: #f1f5f9; padding: 12px 16px; text-align: left; font-size: 13px; font-weight: 600; color: #475569; }}
td {{ padding: 10px 16px; border-top: 1px solid #e2e8f0; font-size: 13px; }}
tr:hover td {{ background: #f8fafc; }}
.footer {{ margin-top: 32px; padding-top: 16px; border-top: 1px solid #e2e8f0; font-size: 12px; color: #94a3b8; text-align: center; }}
</style>
</head>
<body>
<div class="container">
    <div class="header">
        <h1>📋 {task_name}</h1>
        <p>Acdante ITOps Inspection Platform | 生成时间: {timestamp}</p>
    </div>
    <div class="stats">
        <div class="stat stat-ok"><h3>{ok}</h3><p>正常</p></div>
        <div class="stat stat-warn"><h3>{warn}</h3><p>警告</p></div>
        <div class="stat stat-crit"><h3>{crit}</h3><p>严重</p></div>
        <div class="stat stat-score"><h3>{health_score}</h3><p>健康度</p></div>
    </div>
    <h2>巡检摘要</h2>
    <p>{summary}</p>
    <h2>巡检详情</h2>
    <table>
        <thead><tr><th>巡检对象</th><th>巡检项</th><th>分类</th><th>状态</th><th>采集值</th><th>建议</th></tr></thead>
        <tbody>{rows}</tbody>
    </table>
    <div class="footer">Acdante ITOps Inspection Platform v3.1.0 | Powered by Acdante AI</div>
</div>
</body>
</html>"""

    with open(path, "w", encoding="utf-8") as f:
        f.write(html)


def _generate_docx(path, task_name, timestamp, health_score,
                   total, ok, warn, crit, targets, results, summary):
    """生成DOCX报告"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 标题
        title = doc.add_heading(task_name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 信息
        doc.add_paragraph(f"生成时间: {timestamp}")
        doc.add_paragraph(f"健康分数: {health_score}/100")
        doc.add_paragraph(f"巡检摘要: {summary}")

        # 统计表格
        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        headers = ["正常", "警告", "严重", "总项数"]
        values = [str(ok), str(warn), str(crit), str(total)]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for i, v in enumerate(values):
            table.rows[1].cells[i].text = v

        # 详情表格
        doc.add_heading("巡检详情", level=1)
        detail_table = doc.add_table(rows=1 + len(results), cols=5)
        detail_table.style = 'Table Grid'
        detail_headers = ["巡检对象", "巡检项", "分类", "状态", "采集值"]
        for i, h in enumerate(detail_headers):
            detail_table.rows[0].cells[i].text = h

        for idx, r in enumerate(results):
            row = detail_table.rows[idx + 1]
            row.cells[0].text = r.get("target_name", "")
            row.cells[1].text = r.get("item_name", "")
            row.cells[2].text = r.get("category", "")
            row.cells[3].text = r.get("status", "")
            row.cells[4].text = str(r.get("raw_value", ""))[:100]

        doc.save(path)

    except ImportError:
        logger.warning("python-docx未安装，跳过DOCX生成")
        raise


def generate_sample_reports():
    """生成示例报告"""
    sample_results = [
        {"target_name": "APP-SERVER-01", "item_name": "CPU使用率", "category": "CPU", "status": "ok", "raw_value": "45.2%", "suggestion": ""},
        {"target_name": "APP-SERVER-01", "item_name": "内存使用率", "category": "内存", "status": "warning", "raw_value": "87.3%", "suggestion": "内存使用率偏高"},
        {"target_name": "APP-SERVER-01", "item_name": "磁盘使用率", "category": "磁盘", "status": "critical", "raw_value": "/data 95%", "suggestion": "磁盘空间不足"},
    ]
    return generate_report("示例巡检报告", "sample", [], sample_results, format="all")
